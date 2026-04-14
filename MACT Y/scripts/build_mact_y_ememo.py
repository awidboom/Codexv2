from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document

SKILL_SCRIPTS = Path(r"C:\Users\aaw\.codex\skills\barr-frontend\scripts")
if str(SKILL_SCRIPTS) not in sys.path:
    sys.path.append(str(SKILL_SCRIPTS))

from barr_docx_utils import template_to_docx


def set_content_control_text(docx_path: Path, alias_to_text: dict[str, str]) -> None:
    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    xml = files["word/document.xml"].decode("utf-8", errors="ignore")
    for alias, replacement in alias_to_text.items():
        pattern = (
            rf"(<w:alias w:val=\"{re.escape(alias)}\"/>(?:(?!<w:sdtContent>).)*?)"
            rf"(<w:showingPlcHdr/>)?"
            rf"((?:(?!<w:sdtContent>).)*?<w:sdtContent>.*?<w:t>)(.*?)(</w:t>)"
        )

        def _repl(match: re.Match[str]) -> str:
            prefix = match.group(1)
            content_prefix = match.group(3)
            content_suffix = match.group(5)
            return f"{prefix}{content_prefix}{replacement}{content_suffix}"

        xml, count = re.subn(pattern, _repl, xml, count=1, flags=re.DOTALL)
        if count != 1:
            raise RuntimeError(f"Could not replace content control for alias {alias!r}")

    files["word/document.xml"] = xml.encode("utf-8")

    with zipfile.ZipFile(docx_path, "w") as zout:
        for name, data in files.items():
            zout.writestr(name, data)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def build_body(docx_path: Path) -> None:
    doc = Document(str(docx_path))

    while len(doc.paragraphs) > 5:
        remove_paragraph(doc.paragraphs[5])

    sections = [
        (
            "Purpose",
            [
                (
                    "normal",
                    "This memorandum evaluates the proposed amendments to 40 CFR Part 63, Subpart Y for marine tank vessel loading operations and summarizes the likely impacts for the Avon Wharf and Amorco Wharf loading activities reflected in the current Title V permit record.[FN1][FN2][FN4]",
                ),
            ],
        ),
        (
            "Executive Summary",
            [
                (
                    "bullet",
                    "EPA published the Subpart Y technology review proposal on March 4, 2026. As of March 29, 2026, the action remains proposed, and comments are due April 20, 2026.[FN1]",
                ),
                (
                    "bullet",
                    "The proposal is primarily a compliance, monitoring, and reporting update rather than a rewrite of core applicability thresholds. EPA proposes periodic performance testing, electronic reporting, enhanced flare monitoring, removal of startup, shutdown, and malfunction exemptions, and tighter leak monitoring and repair provisions.[FN1][FN3]",
                ),
                (
                    "bullet",
                    "Avon Wharf appears to be the primary directly affected facility because the current permit identifies Source S1560/A1560 as a marine bulk plant with a vapor recovery system and Title V applicable requirements under 40 CFR Part 63 Subpart Y, plus BAAQMD Condition 26406 requiring operation when loading regulated material and 98.5 percent VOC control.[FN2]",
                ),
                (
                    "bullet",
                    "Amorco Wharf appears to be lower impact on the current permit record because the identified loading activity is limited to renewable diesel and is expected to remain well below the 1.5 psia applicability threshold. The permit excerpt does not show a current Subpart Y vapor recovery system or Subpart Y applicable requirements comparable to Avon.[FN4]",
                ),
                (
                    "bullet",
                    "For Avon, the likely compliance impact is moderate and mostly administrative and testing related because Avon already complies through a vapor recovery and fuel-gas routing configuration. The backup flare path is not the primary Subpart Y concern here because refinery flare requirements already address that equipment. For Amorco, direct impact appears limited under the stated renewable diesel-only operating case.[FN2][FN3][FN4]",
                ),
            ],
        ),
        (
            "Detailed Proposed Changes",
            [
                (
                    "bullet",
                    "Remove startup, shutdown, and malfunction exemptions and related maintenance and affirmative defense provisions. EPA's direction is that Subpart Y standards must apply continuously, with comment requested on whether separate maintenance work practices are needed.[FN1][FN3]",
                ),
                (
                    "bullet",
                    "Require electronic reporting through CDX/CEDRI for notifications and reports, with ERT-based submittals for performance tests and CEMS performance evaluations.[FN3]",
                ),
                (
                    "bullet",
                    "Require periodic performance testing at least once every 60 calendar months for control devices currently subject only to an initial performance test.[FN1][FN3]",
                ),
                (
                    "bullet",
                    "Require annual relative accuracy test audits for VOC CEMS where VOC CEMS are used as the compliance monitoring approach.[FN1][FN3]",
                ),
                (
                    "bullet",
                    "Apply enhanced flare monitoring provisions to marine loading flares. EPA's proposal is patterned after refinery flare requirements and could require additional instrumentation, demonstrations, and ongoing monitoring at any Subpart Y flare-based control device.[FN1][FN3]",
                ),
                (
                    "bullet",
                    "Clarify that annual Method 21 monitoring for vapor collection systems and control devices must be performed during active loading of regulated materials subject to 40 CFR 63.562(b), (c), or (d).[FN1][FN3]",
                ),
                (
                    "bullet",
                    "Tighten leak repair requirements so repairs must be completed within 15 days of identification or prior to the next loading operation, whichever is later, and verified by re-monitoring.[FN1][FN3]",
                ),
                (
                    "bullet",
                    "EPA discussed the existing 1.5 psia trigger and other applicability clarifications, but the proposal reviewed here does not appear to expand core Subpart Y applicability to lower-vapor-pressure diesel or renewable diesel loading.[FN1][FN3]",
                ),
            ],
        ),
        (
            "Avon Wharf",
            [
                (
                    "normal",
                    "The current Title V permit identifies Avon Wharf Berth No. 1A, Source S1560, as a marine bulk plant with A1560 vapor recovery. The permit lists loading of gasoline, blendstocks, diesel, distillate, residual oil, renewable diesel, and renewable naphtha, and the applicable requirements table includes 40 CFR Part 63 Subpart Y for S1560/A1560. Under normal operation, the controlled vapors are routed through the vapor recovery system and then into the refinery's 100# fuel gas system. The permit also explains that flare service is a non-routine backup path rather than the normal MACT compliance path.[FN2]",
                ),
                (
                    "bullet",
                    "Direct applicability is clear because Avon already has an identified Subpart Y source and control device in the permit record, along with incorporated Subpart Y citations for MACT applicability, the requirement for a vapor collection system, ship-to-shore compatibility, vapor tightness, and control efficiency.[FN2]",
                ),
                (
                    "bullet",
                    "Periodic performance testing is likely the most concrete new federal obligation for Avon because the proposal would add a recurring test cycle to a source category that historically relied on an initial test plus operating and recordkeeping requirements.[FN1][FN3]",
                ),
                (
                    "bullet",
                    "Electronic reporting would likely require process changes even if the technical control configuration stays the same. Test reports, recurring reports, and notifications would need to move into EPA's CDX/CEDRI workflow.[FN3]",
                ),
                (
                    "bullet",
                    "The proposed active-loading Method 21 clarification and 15-day repair deadline would directly affect Avon if the vapor collection system or control device has leaking components. This is a tighter and more explicit federal compliance expectation than a simple first-attempt repair standard.[FN1][FN3]",
                ),
                (
                    "bullet",
                    "The permit already requires A1560 to operate when loading regulated material and identifies a 98.5 percent VOC control expectation under BAAQMD Condition 26406. The permit also includes local throughput, calculation, pressure recorder/controller, and pressure relief valve monitoring requirements for the Avon system. The proposal therefore looks more like a tightening of federal compliance demonstrations than a requirement to install a fundamentally new control system at Avon.[FN2]",
                ),
                (
                    "bullet",
                    "The reviewed permit language shows Avon primarily routed to vapor recovery and fuel gas handling rather than a dedicated Subpart Y flare. Because the flare system is a non-routine backup path and refinery flare requirements already apply to that equipment, the proposed Subpart Y flare revisions do not appear to be the principal Avon issue.[FN2][FN3]",
                ),
                (
                    "bullet",
                    "Removal of SSM exemptions could increase compliance risk during startup, shutdown, maintenance, or upset conditions affecting the vapor recovery and control system because EPA is moving toward continuous enforceability rather than special treatment for those periods.[FN1][FN3]",
                ),
                (
                    "bullet",
                    "There are already local Bay Area leak and Method 21 requirements relevant to Avon. Rule 8-44 requires in-operation inspections for marine terminal equipment, imposes a 1,000 ppm gaseous leak limit on terminal-side equipment, requires leaks to be minimized within 4 hours, and requires repair before the next controlled operation. Rule 8-44-603 specifies EPA Method 21 for these leak determinations. Separately, the Avon permit adds a six-month Method 21 inspection for the vapor recovery system pressure relief valve with a tighter 500 ppm leak threshold under Condition 26406 Part 8.[FN2][FN6]",
                ),
                (
                    "bullet",
                    "Testing is not starting from zero for a fuel-gas-routed system. BAAQMD Rule 8-44-601 already requires determination of the loading control efficiency during the final 50 percent of the loading event using BAAQMD ST-34, EPA Method 25, EPA Method 25A, or an approved alternative. In practical terms, that is a source test of the marine loading control configuration as operated, even when the normal destination for recovered vapors is the refinery fuel gas system rather than a standalone oxidizer stack.[FN3][FN6]",
                ),
                (
                    "bullet",
                    "Because Avon's normal compliance path is a vapor recovery and fuel-gas routing system, the most defensible reading is that any future periodic test would still be performed as a marine-loading system test rather than as a refinery flare test. That said, the exact test protocol and sampling locations would need to be confirmed against the final rule text and any EPA testing guidance for recovery systems routed to fuel gas. This is an inference from the current local test structure and the proposal's treatment of vapor recovery devices.[FN3][FN6]",
                ),
            ],
        ),
        (
            "Amorco Wharf",
            [
                (
                    "normal",
                    "The permit record reviewed for Plant B2759 is materially different from Avon. The wharf-specific permitted sources identified are the Amorco Wharf slop tank and emergency diesel equipment, and the only loading activity expressly called out is Source S55, 'Amorco Terminal, Renewable Diesel Loading Only,' listed as exempt from permitting under BAAQMD Regulation 2-1-123.3.2 with a firm throughput limit.[FN4]",
                ),
                (
                    "bullet",
                    "On the permit record reviewed here, there is no identified Amorco source equivalent to Avon S1560/A1560 and no clear Subpart Y applicable requirements table tied to a marine vapor recovery system.[FN4]",
                ),
                (
                    "bullet",
                    "Because Amorco is limited to renewable diesel and is expected to remain below 1.5 psia, the current proposal does not appear to create a new direct Subpart Y control requirement for the operating case you described.[FN1][FN3]",
                ),
                (
                    "bullet",
                    "The main remaining Amorco risk is factual rather than regulatory: if the terminal's commodity slate changes to include material above 1.5 psia, or if an unreviewed marine loading control system is put in service, then Amorco could move closer to the Avon-style impact profile and would need a separate applicability review.[FN4]",
                ),
                (
                    "bullet",
                    "Based on the materials reviewed and your operating assumptions, Amorco's near-term implications are best characterized as low direct impact.[FN4]",
                ),
            ],
        ),
        (
            "Recommended Next Steps",
            [
                (
                    "bullet",
                    "Confirm the current commodity slate at Avon and Amorco, with emphasis on keeping the Avon analysis tied to regulated-material loading and the Amorco analysis tied to renewable diesel below 1.5 psia.[FN2][FN4]",
                ),
                (
                    "bullet",
                    "Map each marine-loading control device and compliance monitoring approach now in use, including whether Avon uses VOC CEMS anywhere in the A1560 compliance path and whether any future periodic test can leverage the existing Rule 8-44 loading-event test framework.[FN2][FN3][FN6]",
                ),
                (
                    "bullet",
                    "Compare existing leak monitoring, repair, and reporting practices against the proposal so implementation gaps can be identified before rule finalization, especially where Rule 8-44 already requires Method 21 inspections and repair before the next controlled operation but the federal proposal would extend repair timing to 15 days or before the next loading event, whichever is later.[FN1][FN3][FN6]",
                ),
                (
                    "bullet",
                    "If client comments are being considered, focus comments on how EPA intends periodic performance testing to work for vapor recovery systems routed to refinery fuel gas systems and on how federal leak repair language should interact with existing Rule 8-44 leak inspection and repair requirements.[FN1][FN3][FN6]",
                ),
            ],
        ),
    ]

    for heading, items in sections:
        doc.add_paragraph(heading, style="Heading 1")
        for kind, text in items:
            if kind == "normal":
                doc.add_paragraph(text, style="Normal")
            elif kind == "bullet":
                doc.add_paragraph(text, style="Bullets level 1")
            else:
                raise ValueError(kind)

    doc.save(str(docx_path))


def replace_markers_and_append_references(
    docx_path: Path, marker_map: dict[str, str], references: list[tuple[str, str]]
) -> None:
    doc = Document(str(docx_path))

    for paragraph in doc.paragraphs:
        text = paragraph.text
        for old, new in marker_map.items():
            text = text.replace(old, new)
        if text != paragraph.text:
            paragraph.text = text

    doc.add_paragraph("References", style="Heading 1")
    for label, citation in references:
        doc.add_paragraph(f"{label} {citation}", style="Normal")

    doc.save(str(docx_path))


def main() -> None:
    template_path = Path(r"C:\Users\aaw\.codex\skills\barr-frontend\assets\document-templates\e-memo - Barr.dotx")
    output_dir = Path(r"C:\Users\aaw\Codex\MACT Y\output\doc")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "MACT Y Proposed Rule Barr e-memo updated.docx"

    template_to_docx(template_path, output_path)
    set_content_control_text(
        output_path,
        {
            "Memo type": "Technical Memorandum",
            "to": "Project File",
            "from": "Barr Engineering Co.",
            "subject": "Proposed MACT Y Technology Review Impacts for Avon Wharf and Amorco Wharf",
            "date": "March 29, 2026",
            "project": "MACT Y Rule Review",
        },
    )
    build_body(output_path)
    replace_markers_and_append_references(
        output_path,
        {
            "[FN1]": "[1]",
            "[FN2]": "[2]",
            "[FN3]": "[3]",
            "[FN4]": "[4]",
            "[FN5]": "[5]",
            "[FN6]": "[6]",
        },
        [
            (
                "[1]",
                "EPA, National Emission Standards for Marine Tank Vessel Loading Operations: "
                "Technology Review, proposed rule, 91 Fed. Reg. 10560 (Mar. 4, 2026), Docket "
                "EPA-HQ-OAR-2025-0207.",
            ),
            (
                "[2]",
                "Bay Area AQMD, Final Revision 6 Major Facility Review Permit for Facility B2758 "
                "and B2759 (Sept. 29, 2023); see "
                "Table II A1A and Section IV entries for S1560/A1560 and Condition 26406.",
            ),
            (
                "[3]",
                "EPA-HQ-OAR-2025-0207-0051, Technology Review for National Emission Standards "
                "for Marine Tank Vessel Loading Operations, including sections on electronic "
                "reporting, flare monitoring, performance testing, and monitoring clarifications.",
            ),
            (
                "[4]",
                "Bay Area AQMD, Final Revision 6 Major Facility Review Permit for Facility B2758 "
                "and B2759 (Sept. 29, 2023); see "
                "Table II A2 for B2759 permitted sources and Table II C2 for Source S55, "
                "'Amorco Terminal, Renewable Diesel Loading Only.'",
            ),
            (
                "[5]",
                "EPA-HQ-OAR-2025-0207-0001, proposed Subpart Y preamble discussion of offshore "
                "terminals, active-loading leak monitoring, and clarification that RACT does not "
                "apply offshore.",
            ),
            (
                "[6]",
                "Bay Area AQMD Regulation 8, Rule 44 - Organic Compounds, Marine Vessel Loading "
                "Terminals (Nov. 3, 2021), Sections 8-44-305, 8-44-501, 8-44-601, and 8-44-603.",
            ),
        ],
    )
    print(output_path)


if __name__ == "__main__":
    main()
