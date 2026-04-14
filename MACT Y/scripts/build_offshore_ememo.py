from __future__ import annotations

import sys
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches

SKILL_SCRIPTS = Path(r"C:\Users\aaw\.codex\skills\barr-frontend\scripts")
if str(SKILL_SCRIPTS) not in sys.path:
    sys.path.append(str(SKILL_SCRIPTS))

from barr_docx_utils import replace_markers_with_word_footnotes, template_to_docx


def set_content_control_text(docx_path: Path, alias_to_text: dict[str, str]) -> None:
    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    xml = files["word/document.xml"].decode("utf-8", errors="ignore")
    for alias, replacement in alias_to_text.items():
        pattern = (
            rf"(<w:alias w:val=\"{re.escape(alias)}\"/>(?:(?!<w:sdtContent>).)*?)"
            rf"(<w:showingPlcHdr/>)?"
            rf"((?:(?!<w:sdtContent>).)*?<w:sdtContent>.*?<w:t>)(.*?)(</w:t>)"
        )

        def _repl(match: re.Match[str]) -> str:
            prefix = match.group(1)
            content_prefix = match.group(3)
            content_suffix = match.group(5)
            return f"{prefix}{content_prefix}{replacement}{content_suffix}"

        xml, count = re.subn(pattern, _repl, xml, count=1, flags=re.DOTALL)
        if count != 1:
            raise RuntimeError(f"Could not replace content control for alias {alias!r}")

    files["word/document.xml"] = xml.encode("utf-8")

    with zipfile.ZipFile(docx_path, "w") as zout:
        for name, data in files.items():
            zout.writestr(name, data)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def build_body(docx_path: Path) -> None:
    doc = Document(str(docx_path))

    while len(doc.paragraphs) > 5:
        remove_paragraph(doc.paragraphs[5])

    sections = [
        (
            "Purpose",
            [
                (
                    "normal",
                    "This memorandum summarizes the regulatory history for offshore marine tank vessel loading operations under 40 CFR 63 Subpart Y and reviews whether the 2025 proposed rule reexamined whether add-on controls should be required for loading berths more than 0.5 miles from shore.[FN1]",
                )
            ],
        ),
        (
            "Key Conclusions",
            [
                (
                    "bullet",
                    "Under the current rule, existing offshore loading terminals are not required to install the same MACT or RACT control devices that apply to onshore loading terminals.[FN2]",
                ),
                (
                    "bullet",
                    "EPA historically treated offshore loading differently because offshore vapor control was much more expensive than onshore control, had poor cost effectiveness, and raised technical and environmental concerns.[FN3]",
                ),
                (
                    "bullet",
                    "The 2025 proposal reviews the offshore issue again, but the proposal is framed as a clarification and reaffirmation of the existing offshore treatment rather than a proposal to newly require offshore controls.[FN4]",
                ),
            ],
        ),
        (
            "History Under the Current Rule",
            [
                (
                    "normal",
                    "In the original Subpart Y rulemaking, EPA established separate offshore-terminal and Valdez subcategories. EPA explained that the MACT floor for existing offshore sources was no control of HAP emissions, while the MACT floor for new offshore sources was 95 percent reduction.[FN3]",
                ),
                (
                    "quote",
                    '"For existing offshore terminals, the MACT floor is no control of HAP emissions. For new offshore terminals, the MACT floor is 95 percent reduction of HAP emissions." [FN3]',
                ),
                (
                    "normal",
                    "EPA also concluded that offshore controls would cost roughly two to five times more than onshore controls and would not be reasonable as RACT because of the combination of cost, cost effectiveness, and offshore implementation challenges.[FN3]",
                ),
                (
                    "normal",
                    "The rule also excluded offshore lightering and vessel-to-vessel bulk liquid transfer operations from the source category. Existing offshore terminals instead remained subject to submerged-fill requirements tied to Coast Guard rules.[FN2]",
                ),
                (
                    "normal",
                    "Later amendments retained that structure. The 2011 residual risk and technology review kept the offshore subcategory outside the MACT control-device standards in 40 CFR 63.562(b) and relied on submerged loading requirements for existing offshore sources and certain lower-emitting sources.[FN1]",
                ),
            ],
        ),
        (
            "2025 Proposed Rule Review",
            [
                (
                    "normal",
                    "The 2025 proposed technology review expressly discusses the offshore definition, which remains tied to loading terminals with all loading berths at least 0.81 kilometer, or 0.5 miles, from shore.[FN1]",
                ),
                (
                    "normal",
                    "EPA did prepare new cost analyses for the 2025 proposal, but those analyses address the newly proposed testing, monitoring, reporting, and flare requirements for the source category as a whole rather than a new offshore vapor-control requirement. The economic impact analysis and technology review quantify the cost of periodic performance testing, enhanced flare monitoring, and related compliance measures, but they do not present a separate cost case for requiring existing offshore terminals to install new add-on vapor controls.[FN4][FN5]",
                ),
                (
                    "normal",
                    "The assumptions used in the new 2025 cost work include approximately 190 subject facilities, about 200 control devices, and model terminals derived from 2022 Army Corps loading data. For flare monitoring, EPA assumed 50 flares, with 40 able to use a one-time heating-value demonstration, 8 requiring continuous net-heating-value monitoring, and 2 requiring both heating-value and flow monitoring. EPA also assumed a 7.5 percent interest rate, a 5-year monitoring-equipment life, and 500 hours per year of low-flow supplemental fuel use.[FN4][FN5]",
                ),
                (
                    "normal",
                    "For the monitoring cost-effectiveness analysis itself, EPA used annualized emission reductions of 180 tons per year of HAP and 2,250 tons per year of VOC for periodic performance testing, 100 tons per year of HAP and 1,250 tons per year of VOC for enhanced flare monitoring, and 280 tons per year of HAP and 3,500 tons per year of VOC in total.[FN4]",
                ),
                (
                    "normal",
                    "Those assumptions matter for the proposed monitoring revisions, but they do not appear to have been used to justify retaining the offshore exemption. Instead, the technology review states that EPA identified no new loading technologies, that submerged loading is already required by Coast Guard regulations, and that EPA only evaluated whether revisions to the 10/25-ton threshold and the 1.5 psia exemption would produce cost-effective reductions. The memorandum does not identify a comparable new cost analysis of the offshore exemption in 40 CFR 63.560(d)(6).[FN4]",
                ),
                (
                    "normal",
                    "Consistent with that approach, the proposal text focuses on clarifying that 40 CFR 63.560(d)(6) excludes existing offshore loading terminals from the MACT standards in 40 CFR 63.562(b), and EPA further proposes to clarify that the RACT provisions do not apply offshore. The redline confirms that existing offshore terminals would continue to meet submerged-fill requirements under 46 CFR 153.282 rather than a broader vapor control requirement, and the supporting technology review tables continue to show submerged loading as the operative requirement for existing offshore categories.[FN1][FN2][FN4]",
                ),
            ],
        ),
        (
            "Why Controls Were Not Required Again",
            [
                (
                    "normal",
                    "The best reading of the 2025 record is that EPA did not perform a fresh offshore-specific cost-effectiveness analysis because the technology review did not identify a new control technology or other development that would justify reopening the original offshore determination. EPA instead relied on the historical record showing offshore controls were much more expensive and less cost-effective than onshore controls, together with the current finding that no new loading technologies were identified in the federal, state, local, RBLC, and international review.[FN3][FN4]",
                ),
                (
                    "normal",
                    "EPA also points to practical feasibility considerations that continue to distinguish offshore operations. The 2025 technology review notes that some offshore facilities use flexible or submerged vapor lines for which Method 21 monitoring is infeasible, and it proposes allowing offshore facilities to rely on audible, visual, or olfactory leak detection in those circumstances rather than imposing more burdensome instrument-monitoring requirements.[FN4]",
                ),
                (
                    "normal",
                    "On future rulemaking, the proposal does not say that offshore controls are likely to be required later. The most forward-looking discussion is a specific request for comment asking how MACT and RACT standards should apply to offshore marine terminals. That keeps the offshore issue open in this rulemaking, but the current record reviewed here does not identify a likely future path to mandatory offshore vapor controls absent new technical, cost, or safety information.[FN1]",
                ),
                (
                    "normal",
                    "Accordingly, the most defensible explanation is that EPA reviewed offshore loading again, preserved the existing offshore structure, and proposed only clarifying changes. Existing offshore loading terminals remain outside the main MACT and RACT control-device requirements, while new offshore major sources remain subject to the 95 percent reduction standard identified in the original rulemaking.[FN3][FN4]",
                ),
            ],
        ),
    ]

    for heading, items in sections:
        doc.add_paragraph(heading, style="Heading 1")
        for kind, text in items:
            if kind == "normal":
                doc.add_paragraph(text, style="Normal")
            elif kind == "bullet":
                doc.add_paragraph(text, style="Bullets level 1")
            elif kind == "quote":
                paragraph = doc.add_paragraph(style="Normal")
                paragraph.paragraph_format.left_indent = Inches(0.5)
                run = paragraph.add_run(text)
                run.italic = True
            else:
                raise ValueError(kind)

    doc.save(str(docx_path))


def main() -> None:
    template_path = Path(r"C:\Users\aaw\.codex\skills\barr-frontend\assets\document-templates\e-memo - Barr.dotx")
    output_path = Path(r"C:\Users\aaw\Codex\MACT Y\Subpart Y Offshore Loading e-memo revised 2.docx")

    template_to_docx(template_path, output_path)
    set_content_control_text(
        output_path,
        {
            "Memo type": "Technical Memorandum",
            "to": "Project File",
            "from": "Regulatory Analysis Support",
            "subject": "Offshore Loading Under 40 CFR 63 Subpart Y",
            "date": "March 29, 2026",
            "project": "MACT Y Rule History Review",
        },
    )
    build_body(output_path)
    replace_markers_with_word_footnotes(
        output_path,
        {
            "[FN1]": (
                "EPA-HQ-OAR-2025-0207-0001, National Emission Standards for Marine Tank Vessel "
                "Loading Operations: Technology Review, pp. 5-7, 12, 18."
            ),
            "[FN2]": (
                "EPA-HQ-OAR-2025-0207-0010, Redline of 40 CFR Part 63 Subpart Y Showing Proposed "
                "Amendments, pp. 2-3."
            ),
            "[FN3]": (
                "EPA-HQ-OAR-2025-0207-0030, NESHAP Subpart Y Technical Support for Final Standards: "
                "Summary of Public Comments and Responses, pp. 13, 21-22, 93, 95-96."
            ),
            "[FN4]": (
                "EPA-HQ-OAR-2025-0207-0051, Technology Review for National Emission Standards for Marine "
                "Tank Vessel Loading Operations, pp. 5, 10-11, 17, 24-26."
            ),
            "[FN5]": (
                "EPA-HQ-OAR-2025-0207-0017, Economic Impact Analysis, pp. 5-6."
            ),
        },
    )
    print(output_path)


if __name__ == "__main__":
    main()
