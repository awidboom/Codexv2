from __future__ import annotations

from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"C:\Users\aaw\.codex\skills\barr-frontend\assets\document-templates\e-memo - Barr.dotx")
OUTPUT = ROOT / "outputs" / "Ecobat_NOV_financial_obligation_e-memo.docx"


SECTIONS = [
    (
        "Settlement status",
        [
            "Has the June 2025 global settlement letter been accepted, rejected, countered, or superseded?",
            "Was a settlement agreement for P79157, P79172, P76003, P79476, and P68909 ever executed?",
            "If not executed, what is the current negotiation status, latest agency communication date, and Ecobat's current counterposition?",
            "Has SCAQMD updated its penalty demand since the letter's stated total settlement amount?",
            "Has CPI or any other escalation been applied since that offer?",
            "Are there any side letters, tolling agreements, payment plans, or informal settlement understandings not included in the current files?",
        ],
    ),
    (
        "Notice-by-notice exposure",
        [
            "For each NOV, what amount does Ecobat currently believe is still unpaid, unresolved, or subject to negotiation?",
            "Which NOVs does Ecobat believe are fully closed as a legal matter, and what evidence supports closure?",
            "Which NOVs remain open administratively even if Ecobat believes compliance has been achieved?",
            "For P80830, P69321, and P69327, has SCAQMD made any formal or informal penalty demand?",
            "For older historical NOVs, including P67058, are there any unresolved obligations, reopening risk, or reserved rights?",
        ],
    ),
    (
        "Accounting and reserves",
        [
            "What reserve, accrual, or contingent liability has Ecobat booked for AQMD NOVs as of the most recent month-end?",
            "How has that reserve changed over the last 12 months?",
            "What assumptions underlie the reserve by NOV?",
            "Has outside counsel provided a loss contingency range, and if so what is the low, expected, and high case?",
            "Are any NOV-related amounts classified as probable, reasonably possible, or remote under Ecobat's accounting framework?",
            "Are any penalties expected to be capitalized, expensed, indemnified, or insured?",
        ],
    ),
    (
        "Penalty drivers",
        [
            "What specific mitigation arguments is Ecobat relying on for each open NOV?",
            "What documentation has Ecobat provided to support those arguments?",
            "Which penalty components does Ecobat concede versus still dispute?",
            "Has Ecobat modeled best-case, likely-case, and worst-case settlement outcomes?",
            "Does Ecobat expect SCAQMD to seek only civil penalties, or also injunctive or operational commitments with cost implications?",
        ],
    ),
    (
        "Compliance closure",
        [
            "What corrective actions has Ecobat completed for each NOV?",
            "What corrective actions remain outstanding, and what will they cost?",
            "Have any permit modifications been submitted, approved, or denied that affect the alleged violations?",
            "Has SCAQMD confirmed any violations as cured or accepted compliance only for any items beyond what appears in the current settlement correspondence?",
            "Are there any required monitoring upgrades, source tests, reporting changes, staffing changes, or consultant work still needed to close the matters?",
        ],
    ),
    (
        "Future operating and capital cost",
        [
            "What recurring cost is expected to maintain compliance going forward for pressure monitoring, pH control, source testing, QCER review, and reporting?",
            "What capital projects are still needed to reduce repeat NOV risk?",
            "What is the cost and schedule for those projects?",
            "Are there any expected downtime, production limits, or operational inefficiencies tied to compliance commitments?",
            "Has Ecobat quantified the cost of additional ambient monitoring, data validation, or third-party auditing?",
        ],
    ),
    (
        "Agency process risk",
        [
            "What is the current SCAQMD enforcement contact and status for each open NOV?",
            "Are any matters at risk of referral to litigation?",
            "Has Ecobat missed any response deadlines, payment deadlines, or document-production requests?",
            "Has SCAQMD indicated that penalties could increase if negotiations continue?",
            "Are there any parallel EPA or federal implications tied to the AQMD matters?",
        ],
    ),
    (
        "Supporting documents to request",
        [
            "Latest NOV tracker with status, owner, expected closure date, and estimated liability.",
            "All settlement communications after June 2025.",
            "Any executed settlement agreements and proof of payment.",
            "Internal reserve memo or loss contingency analysis.",
            "Corrective action plans and evidence of completion.",
            "Any board, audit, or management summaries discussing AQMD exposure.",
        ],
    ),
]


def template_to_docx_bytes(path: Path) -> bytes:
    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
        tmp.write(path.read_bytes())
    try:
        with ZipFile(tmp_path, "r") as zf:
            members = {name: zf.read(name) for name in zf.namelist()}
            xml = members["[Content_Types].xml"].decode("utf-8")
            xml = xml.replace(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
            )
            members["[Content_Types].xml"] = xml.encode("utf-8")

        with ZipFile(tmp_path, "w", compression=ZIP_DEFLATED) as zf:
            for name, data in members.items():
                zf.writestr(name, data)

        return tmp_path.read_bytes()
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


def set_header_field(paragraph, label: str, value: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(value)


def build_doc() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT.write_bytes(template_to_docx_bytes(TEMPLATE))

    doc = Document(OUTPUT)
    header_values = [
        ("To", "Ecobat management and counsel"),
        ("From", "Barr due diligence team"),
        ("Subject", "Questions to assess future financial obligation for AQMD NOV closure"),
        ("Date", "March 25, 2026"),
        ("Project", "Ecobat due diligence"),
        ("c", ""),
    ]
    for paragraph, (label, value) in zip(doc.paragraphs[:6], header_values):
        set_header_field(paragraph, label, value)

    intro = doc.add_paragraph(style="Memorandum")
    intro.add_run(
        "Purpose. The questions below are intended to establish the current legal, financial, and "
        "operational exposure tied to AQMD notices of violation, including unpaid or unbooked "
        "penalties, unresolved settlement positions, remaining corrective-action costs, and the risk "
        "of additional obligations if negotiations continue or break down."
    )

    for title, questions in SECTIONS:
        doc.add_paragraph(title, style="Heading 1")
        for question in questions:
            doc.add_paragraph(question, style="Bullets level 1")

    closing = doc.add_paragraph(style="Memorandum")
    closing.add_run(
        "Requested support documents should include the latest NOV tracker, post-June 2025 settlement "
        "communications, executed settlement agreements, proof of payment, reserve analyses, "
        "corrective-action evidence, and any board or audit summaries discussing AQMD exposure."
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(f"Wrote {OUTPUT}")
