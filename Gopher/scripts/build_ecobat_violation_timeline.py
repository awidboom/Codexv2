from __future__ import annotations

import csv
from collections import defaultdict
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt

from aqmd_nov_summary_rows import AQMD_NOV_SUMMARY_HEADERS, AQMD_NOV_SUMMARY_ROWS


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs"
OUTPUT_PATH = OUTPUT_DIR / "Ecobat_violation_timeline.docx"
NOV_CSV = ROOT / "aqmd_find_nov" / "open_novs" / "open_novs_index.csv"


@dataclass(frozen=True)
class SourceRef:
    label: str
    page: int


@dataclass(frozen=True)
class GeckoType:
    key: str
    title: str
    requirement: str
    equipment: str
    classification: str
    summary: str
    source: SourceRef
    event_times: tuple[str, ...]
    cause: str
    corrective_action: str


GECKO_TYPES: tuple[GeckoType, ...] = (
    GeckoType(
        key="c169_np",
        title="Battery Wrecker enclosure negative pressure not maintained",
        requirement="Rule 1420.1(e)(3)",
        equipment="C169 / Battery Wrecker enclosure",
        classification="Self-disclosed deviation in Title V SAM/ACC",
        summary="Negative pressure excursions in the Battery Wrecker enclosure.",
        source=SourceRef(
            "mar23/Gecko_45_Ecobat Res California (8547) Title V SAM_010125-063025 (signed).pdf",
            6,
        ),
        event_times=(
            "2025-06-25 08:15",
            "2025-06-25 08:30",
            "2025-06-25 08:45",
            "2025-11-12 06:40",
            "2025-12-05 04:15",
            "2025-12-05 04:30",
            "2025-12-05 04:45",
            "2025-12-05 05:00",
            "2025-12-05 05:15",
            "2025-12-05 05:30",
            "2025-12-05 05:45",
            "2025-12-05 06:00",
            "2025-12-05 06:15",
            "2025-12-05 06:30",
        ),
        cause=(
            "June event tied to plant power shutdown for new CEMS shelter work; later events tied to Edison "
            "power shutdown / PLC fault reset."
        ),
        corrective_action=(
            "Require mitigation plans for power work affecting Busch Units; PLC output card replaced and alarms/text "
            "alerts added."
        ),
    ),
    GeckoType(
        key="d8_d84_np",
        title="Furnace negative pressure not maintained",
        requirement="Rule 1420.1(f)(3)",
        equipment="D8 and D84",
        classification="Self-disclosed deviation in Title V SAM/ACC",
        summary="Insufficient negative pressure for the furnace containment areas.",
        source=SourceRef(
            "mar23/Gecko_45_Ecobat Res California (8547) Title V SAM_070125-123125 (signed).pdf",
            9,
        ),
        event_times=(
            "2025-09-12 19:30",
            "2025-09-12 20:00",
        ),
        cause="SCADA and RKI systems were down after battery-backup / switch faults.",
        corrective_action="Switched output mode on, cleared faults, and restored monitoring/output.",
    ),
    GeckoType(
        key="c89_flow",
        title="EAF scrubber recirculation flow below permit limit",
        requirement="Permit Condition C8.3",
        equipment="C89",
        classification="Self-disclosed deviation in Title V SAM/ACC",
        summary="EAF primary and/or secondary scrubber flow fell below the 66 gpm permit limit.",
        source=SourceRef(
            "mar23/Gecko_45_Ecobat Res California (8547) Title V SAM_010125-063025 (signed).pdf",
            10,
        ),
        event_times=(
            "2025-01-20 21:00",
            "2025-01-26 10:00",
            "2025-01-26 11:00",
            "2025-02-01 09:00",
            "2025-02-01 13:00",
            "2025-02-01 20:00",
            "2025-02-01 21:00",
            "2025-03-12 21:00",
            "2025-03-14 18:00",
            "2025-03-14 19:00",
            "2025-03-17 09:00",
            "2025-03-18 08:00",
            "2025-03-18 09:00",
            "2025-03-18 10:00",
            "2025-03-18 13:00",
            "2025-03-19 15:00",
            "2025-03-31 22:00",
            "2025-04-06 08:00",
            "2025-04-06 11:00",
            "2025-04-06 12:00",
            "2025-04-06 13:00",
            "2025-04-06 14:00",
            "2025-04-09 02:00",
            "2025-04-25 03:00",
            "2025-05-07 19:00",
            "2025-05-10 20:00",
            "2025-05-11 12:00",
            "2025-05-11 17:00",
            "2025-05-13 21:00",
            "2025-05-13 22:00",
            "2025-05-13 23:00",
            "2025-06-29 00:00",
            "2025-06-29 01:00",
            "2025-06-29 02:00",
            "2025-06-29 03:00",
            "2025-06-29 22:00",
            "2025-06-29 23:00",
            "2025-07-06 06:00",
            "2025-07-06 07:00",
            "2025-07-06 08:00",
            "2025-07-06 09:00",
            "2025-07-11 22:00",
            "2025-07-13 11:00",
            "2025-07-19 10:00",
            "2025-07-19 12:00",
            "2025-07-19 13:00",
            "2025-07-19 14:00",
            "2025-08-09 23:00",
            "2025-08-10 00:00",
            "2025-09-12 19:00",
            "2025-09-12 20:00",
            "2025-09-12 21:00",
            "2025-09-12 22:00",
            "2025-09-20 20:00",
            "2025-09-20 21:00",
            "2025-09-20 22:00",
            "2025-09-20 23:00",
            "2025-09-21 00:00",
            "2025-09-21 23:00",
            "2025-09-22 00:00",
            "2025-09-22 01:00",
            "2025-09-22 02:00",
            "2025-09-22 03:00",
            "2025-09-22 04:00",
            "2025-09-22 05:00",
            "2025-09-22 06:00",
            "2025-09-24 00:00",
            "2025-09-24 01:00",
            "2025-09-25 05:00",
            "2025-10-08 09:00",
            "2025-10-08 10:00",
            "2025-10-08 11:00",
            "2025-10-08 12:00",
        ),
        cause="Recirculation rate drifted.",
        corrective_action="Evaluate larger piping and install redundant pump / permit modification.",
    ),
    GeckoType(
        key="c89_ph",
        title="EAF scrubber pH below permit limit",
        requirement="Permit Condition C8.2",
        equipment="C89",
        classification="Self-disclosed deviation in Title V SAM/ACC",
        summary="EAF scrubber pH fell below the listed 9.4 permit limit.",
        source=SourceRef(
            "mar23/Gecko_45_Ecobat Res California (8547) Title V SAM_010125-063025 (signed).pdf",
            14,
        ),
        event_times=(
            "2025-01-02 10:00",
            "2025-01-02 17:00",
            "2025-01-21 06:00",
            "2025-01-21 16:00",
            "2025-01-22 07:00",
            "2025-01-28 17:00",
            "2025-01-28 18:00",
            "2025-01-29 17:00",
            "2025-01-31 06:00",
            "2025-02-03 15:00",
            "2025-02-03 16:00",
            "2025-02-03 17:00",
            "2025-02-03 18:00",
            "2025-02-03 19:00",
            "2025-02-27 01:00",
            "2025-02-27 02:00",
            "2025-03-18 10:00",
            "2025-04-04 15:00",
            "2025-04-06 07:00",
            "2025-04-06 16:00",
            "2025-04-06 17:00",
            "2025-04-06 18:00",
            "2025-04-06 19:00",
            "2025-04-06 20:00",
            "2025-04-06 21:00",
            "2025-04-06 22:00",
            "2025-04-06 23:00",
            "2025-04-07 00:00",
            "2025-04-07 01:00",
            "2025-04-07 02:00",
            "2025-04-07 03:00",
            "2025-04-10 17:00",
            "2025-04-10 18:00",
            "2025-05-03 22:00",
            "2025-05-03 23:00",
            "2025-05-04 00:00",
            "2025-05-04 01:00",
            "2025-05-04 02:00",
            "2025-05-04 03:00",
            "2025-05-04 04:00",
            "2025-05-04 13:00",
            "2025-05-04 14:00",
            "2025-05-04 15:00",
            "2025-05-04 16:00",
            "2025-05-04 17:00",
            "2025-05-04 18:00",
            "2025-05-05 03:00",
            "2025-05-05 04:00",
            "2025-05-05 05:00",
            "2025-05-05 06:00",
            "2025-05-05 07:00",
            "2025-05-06 21:00",
            "2025-05-06 22:00",
            "2025-05-06 23:00",
            "2025-05-07 00:00",
            "2025-05-07 01:00",
            "2025-05-07 15:00",
            "2025-05-07 16:00",
            "2025-05-08 01:00",
            "2025-05-08 02:00",
            "2025-05-08 03:00",
            "2025-05-08 04:00",
            "2025-05-08 05:00",
            "2025-05-08 06:00",
            "2025-05-08 16:00",
            "2025-05-08 17:00",
            "2025-05-10 06:00",
            "2025-05-10 07:00",
            "2025-05-10 08:00",
            "2025-05-10 09:00",
            "2025-05-10 13:00",
            "2025-05-10 14:00",
            "2025-05-10 15:00",
            "2025-05-14 07:00",
            "2025-07-16 09:00",
            "2025-07-23 14:00",
            "2025-07-31 15:00",
            "2025-07-31 16:00",
            "2025-09-22 22:00",
        ),
        cause="pH probe operating in harsh conditions; Ecobat also argues the operative limit should be 8.0, not 9.4.",
        corrective_action="Increased line cleaning and pH probe calibration; permit modifications submitted.",
    ),
    GeckoType(
        key="c21_temp",
        title="Refinery baghouse exhaust temperature not recorded / above limit",
        requirement="Permit Condition C6.3",
        equipment="C21",
        classification="Self-disclosed deviation in Title V ACC",
        summary="Refinery baghouse exhaust temperature was invalid / above permit limit due to data logging issues.",
        source=SourceRef(
            "mar23/Gecko_45_Ecobat Res California (8547) Title V SAM_010125-063025 (signed).pdf",
            19,
        ),
        event_times=(
            "2025-05-13 13:00",
            "2025-05-13 14:00",
            "2025-05-13 17:00",
            "2025-05-13 18:00",
            "2025-05-13 19:00",
            "2025-05-13 20:00",
            "2025-05-13 21:00",
            "2025-05-13 22:00",
            "2025-05-14 00:00",
            "2025-05-14 01:00",
            "2025-05-14 02:00",
            "2025-05-14 03:00",
            "2025-05-14 04:00",
            "2025-05-14 05:00",
            "2025-05-14 06:00",
            "2025-05-14 07:00",
            "2025-05-14 08:00",
        ),
        cause="Data logging issue.",
        corrective_action="Establish redundant data recording for the parameter.",
    ),
    GeckoType(
        key="wesp_voltage",
        title="WESP transformer/rectifier voltage not maintained",
        requirement="Permit Condition D12.16",
        equipment="C141, C145, C149, C153, C157",
        classification="Self-disclosed deviation in Title V SAM/ACC",
        summary="High-voltage electric circuit serving the WESP did not maintain set voltage during an unplanned shutdown.",
        source=SourceRef(
            "mar23/Gecko_45_Ecobat Res California (8547) Title V SAM_070125-123125 (signed).pdf",
            12,
        ),
        event_times=("2025-10-14 12:45",),
        cause="Unplanned shutdown after CB22 tripped, affecting e-stops and control relays.",
        corrective_action="WESP restarted; voltage returned to permit limit; event reported to SCAQMD.",
    ),
)


def parse_ts(value: str) -> datetime:
    return datetime.strptime(value, "%Y-%m-%d %H:%M")


def fmt_date(value: datetime) -> str:
    return value.strftime("%b %d, %Y")


def fmt_dt(value: datetime) -> str:
    return value.strftime("%b %d, %Y %H:%M")


def style_doc(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    for style_name, size in [("Title", 18), ("Heading 1", 14), ("Heading 2", 11)]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)


def load_novs() -> list[dict[str, str]]:
    with NOV_CSV.open(newline="", encoding="utf-8-sig") as handle:
        rows = list(csv.DictReader(handle))
    for row in rows:
        row["violation_dt"] = datetime.strptime(row["violation_date"], "%m/%d/%Y")
        row["issue_dt"] = datetime.strptime(row["issue_date"], "%m/%d/%Y")
    rows.sort(key=lambda row: (row["violation_dt"], row["issue_dt"], row["notice_number"]))
    return rows


def build_gecko_daily_entries() -> list[dict[str, object]]:
    entries: list[dict[str, object]] = []
    for violation in GECKO_TYPES:
        grouped: dict[datetime.date, list[datetime]] = defaultdict(list)
        for stamp in violation.event_times:
            parsed = parse_ts(stamp)
            grouped[parsed.date()].append(parsed)

        for day, stamps in sorted(grouped.items()):
            stamps.sort()
            entries.append(
                {
                    "date": datetime.combine(day, datetime.min.time()),
                    "source_group": "Gecko 45 self-disclosed responses",
                    "type": violation.title,
                    "details": (
                        f"{violation.requirement}; {violation.equipment}; "
                        f"{len(stamps)} timestamp(s) from {fmt_dt(stamps[0])} to {fmt_dt(stamps[-1])}. "
                        f"{violation.summary}"
                    ),
                    "status": violation.classification,
                    "source_ref": f"{violation.source.label} (page {violation.source.page})",
                }
            )
    entries.sort(key=lambda item: (item["date"], item["type"]))
    return entries


def add_title(document: Document) -> None:
    document.add_heading("Ecobat / Facility 8547 Violation Timeline", level=0)
    p = document.add_paragraph()
    p.add_run(
        "Built from the three Gecko 45 Title V response PDFs in "
        "`mar23` and the AQMD NOV index in `aqmd_find_nov/open_novs/open_novs_index.csv`."
    )
    p.add_run().add_break()
    p.add_run(
        "This distinguishes self-disclosed Title V deviations from AQMD-issued NOVs. "
        "For NOV rows, the timeline is sorted by the listed violation date, with issue date shown separately."
    )
    outlook = document.add_paragraph()
    outlook.add_run(
        "Forward-looking note. Based on the June 12, 2025 global settlement letter, SCAQMD appears to be taking a strict-liability, "
        "records-driven approach: it is generally unwilling to excuse violations based on monitor issues, contractor error, or Ecobat's "
        "view that emissions impacts were unlikely, but it may narrow or mitigate penalties where Ecobat provides specific operational context, "
        "documented corrective action, or evidence that a cited requirement should be treated as satisfied. On that pattern, the more recent 2025 "
        "violations and NOVs not covered by the global settlement letter may still face material penalty exposure if SCAQMD concludes the records "
        "show repeat monitoring, reporting, pressure-control, or maintenance deficiencies, though Ecobat may retain leverage to reduce or compartmentalize "
        "those amounts where the events were brief, occurred during shutdowns or rebuilds, or can be tied to prompt reporting and targeted remediation."
    )


def add_gecko_summary(document: Document) -> None:
    document.add_heading("Gecko 45 Violation Types", level=1)
    table = document.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = [
        "Violation Type",
        "Requirement",
        "Equipment",
        "Date Span",
        "Affected Days",
        "Cause",
        "Corrective Action / Position",
        "Source",
    ]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header

    for violation in GECKO_TYPES:
        stamps = sorted(parse_ts(value) for value in violation.event_times)
        days = sorted({stamp.date() for stamp in stamps})
        day_list = ", ".join(day.strftime("%Y-%m-%d") for day in days)
        row = table.add_row().cells
        row[0].text = violation.title
        row[1].text = violation.requirement
        row[2].text = violation.equipment
        row[3].text = f"{fmt_date(stamps[0])} to {fmt_date(stamps[-1])}"
        row[4].text = day_list
        row[5].text = violation.cause
        row[6].text = violation.corrective_action
        row[7].text = f"{violation.source.label} (page {violation.source.page})"


def add_nov_summary(document: Document, nov_rows: list[dict[str, str]]) -> None:
    document.add_heading("AQMD NOV Summary", level=1)
    table = document.add_table(rows=1, cols=len(AQMD_NOV_SUMMARY_HEADERS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, header in zip(table.rows[0].cells, AQMD_NOV_SUMMARY_HEADERS):
        cell.text = header

    for row_data in AQMD_NOV_SUMMARY_ROWS:
        row = table.add_row().cells
        row[0].text = row_data["notice"]
        row[1].text = row_data["violation_date"]
        row[2].text = row_data["issue_date"]
        row[3].text = row_data["status"]
        row[4].text = row_data["rules"]
        row[5].text = row_data["violation_description"]
        row[6].text = row_data["ecobat_response"]
        row[7].text = row_data["scaqmd_response"]
        row[8].text = row_data["offered_penalty"]


def add_combined_timeline(document: Document, nov_rows: list[dict[str, str]]) -> None:
    document.add_page_break()
    document.add_heading("Combined Chronology", level=1)

    entries = build_gecko_daily_entries()
    for nov in nov_rows:
        entries.append(
            {
                "date": nov["violation_dt"],
                "source_group": "AQMD NOV",
                "type": f"{nov['notice_number']} ({nov['follow_up_status'] or 'status not stated'})",
                "details": (
                    f"Violation date {nov['violation_dt'].strftime('%Y-%m-%d')}; issue date {nov['issue_dt'].strftime('%Y-%m-%d')}; "
                    f"{nov['violation_description']}"
                ),
                "status": nov["rules"],
                "source_ref": nov["find_detail_html"],
            }
        )

    entries.sort(key=lambda item: (item["date"], item["source_group"], item["type"]))

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Date", "Source Group", "Violation / Notice", "Details", "Rule / Status", "Source"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header

    for entry in entries:
        row = table.add_row().cells
        row[0].text = entry["date"].strftime("%Y-%m-%d")
        row[1].text = str(entry["source_group"])
        row[2].text = str(entry["type"])
        row[3].text = str(entry["details"])
        row[4].text = str(entry["status"])
        row[5].text = str(entry["source_ref"])


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    nov_rows = load_novs()

    document = Document()
    style_doc(document)
    add_title(document)
    add_gecko_summary(document)
    add_nov_summary(document, nov_rows)
    add_combined_timeline(document, nov_rows)
    document.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
