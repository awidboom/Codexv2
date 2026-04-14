from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "outputs" / "Ecobat_violation_timeline.docx"
OUTLOOK_TEXT = (
    "Forward-looking note. Based on the June 12, 2025 global settlement letter, SCAQMD appears to be taking a strict-liability, "
    "records-driven approach: it is generally unwilling to excuse violations based on monitor issues, contractor error, or Ecobat's "
    "view that emissions impacts were unlikely, but it may narrow or mitigate penalties where Ecobat provides specific operational context, "
    "documented corrective action, or evidence that a cited requirement should be treated as satisfied. On that pattern, the more recent 2025 "
    "violations and NOVs not covered by the global settlement letter may still face material penalty exposure if SCAQMD concludes the records "
    "show repeat monitoring, reporting, pressure-control, or maintenance deficiencies, though Ecobat may retain leverage to reduce or compartmentalize "
    "those amounts where the events were brief, occurred during shutdowns or rebuilds, or can be tied to prompt reporting and targeted remediation."
)


def main() -> None:
    doc = Document(DOCX_PATH)
    if any(p.text.strip() == OUTLOOK_TEXT for p in doc.paragraphs):
        doc.save(DOCX_PATH)
        print(f"Already present in {DOCX_PATH}")
        return

    insert_idx = 2 if len(doc.paragraphs) >= 2 else len(doc.paragraphs)
    paragraph = doc.paragraphs[insert_idx].insert_paragraph_before(OUTLOOK_TEXT)
    paragraph.style = doc.styles["Normal"]
    doc.save(DOCX_PATH)
    print(f"Updated {DOCX_PATH}")


if __name__ == "__main__":
    main()
