import argparse
import datetime as dt
import subprocess
import sys
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from PIL import Image, ImageDraw, ImageFont


DEFAULT_SHEET = "Project Det With Timesheet Com"


@dataclass(frozen=True)
class WeekWindow:
    week_posted: dt.date  # expected to be a Friday
    worked_start: dt.date  # Saturday
    worked_end: dt.date  # Friday


@dataclass(frozen=True)
class WeekExtract:
    df: pd.DataFrame
    total_hours: float
    blank_person_hours: float
    blank_project_hours: float


def parse_date(value: str) -> dt.date:
    try:
        return dt.date.fromisoformat(value)
    except ValueError as exc:
        raise SystemExit(f"Invalid date: {value!r}. Expected YYYY-MM-DD.") from exc


def week_window(week_posted: dt.date) -> WeekWindow:
    # Week Posted is the posting week (Fri), but hours reflect the prior worked week (Sat–Fri),
    # ending 7 days before Week Posted.
    worked_end = week_posted - dt.timedelta(days=7)
    worked_start = worked_end - dt.timedelta(days=6)
    return WeekWindow(week_posted=week_posted, worked_start=worked_start, worked_end=worked_end)


def _clean_text(value) -> str | None:
    if value is None:
        return None
    if isinstance(value, float) and pd.isna(value):
        return None
    text = str(value).strip()
    if not text:
        return None
    if text.lower() in {"(blank)", "nan", "none"}:
        return None
    return text


def detect_header_row(workbook: Path, sheet: str, max_scan_rows: int = 200) -> int:
    raw = pd.read_excel(
        workbook, sheet_name=sheet, engine="openpyxl", header=None, nrows=max_scan_rows
    )
    for row_idx in range(len(raw)):
        row = raw.iloc[row_idx].astype(str)
        if row.str.contains("Week Posted", case=False, na=False).any() and row.str.contains(
            "Project", case=False, na=False
        ).any():
            return row_idx
    raise SystemExit(
        f"Could not detect header row on sheet {sheet!r}. "
        "Try passing --header-row explicitly."
    )


def detect_latest_week_posted(workbook: Path, sheet: str, header_row: int) -> dt.date:
    df = pd.read_excel(
        workbook, sheet_name=sheet, engine="openpyxl", header=header_row, usecols=["Week Posted"]
    )
    series = pd.to_datetime(df["Week Posted"], errors="coerce", format="mixed").dt.date.dropna()
    if series.empty:
        raise SystemExit(f"No 'Week Posted' values found on sheet {sheet!r}.")
    return max(series)


def _workbook_lockfile_path(workbook: Path) -> Path:
    return workbook.with_name(f"~${workbook.name}")


def refresh_workbook_via_excel_com(workbook: Path, timeout_s: int = 600) -> None:
    if sys.platform != "win32":
        raise SystemExit("Auto-refresh is only supported on Windows.")

    try:
        proc_count = subprocess.check_output(
            [
                "powershell",
                "-NoProfile",
                "-Command",
                "(Get-Process EXCEL -ErrorAction SilentlyContinue | Measure-Object).Count",
            ],
            text=True,
        ).strip()
        if proc_count and int(proc_count) > 0:
            raise SystemExit(
                "Excel is currently running. Close all Excel instances before running this script so the "
                "workbook refresh/save can complete."
            )
    except (OSError, ValueError):
        pass

    lockfile = _workbook_lockfile_path(workbook)
    if lockfile.exists():
        raise SystemExit(
            f"Workbook appears to be open (lock file exists): {lockfile.name!r}. "
            "Close Excel and retry, or run with --no-refresh."
        )

    ps = f"""
$ErrorActionPreference = 'Stop'
$path = '{str(workbook.resolve()).replace("'", "''")}'
$excel = New-Object -ComObject Excel.Application
$excel.Visible = $false
$excel.DisplayAlerts = $false
$excel.AskToUpdateLinks = $false
$excel.EnableEvents = $false
try {{
  $wb = $excel.Workbooks.Open($path, 0, $false)
  try {{
    foreach ($c in $wb.Connections) {{
      try {{ $c.RefreshWithRefreshAll = $true }} catch {{ }}
      try {{ $c.OLEDBConnection.BackgroundQuery = $false }} catch {{ }}
      try {{ $c.ODBCConnection.BackgroundQuery = $false }} catch {{ }}
    }}
  }} catch {{ }}
  try {{
    foreach ($ws in $wb.Worksheets) {{
      try {{
        foreach ($qt in $ws.QueryTables) {{ $qt.BackgroundQuery = $false }}
      }} catch {{ }}
    }}
  }} catch {{ }}
  $wb.RefreshAll() | Out-Null
  try {{ $excel.CalculateUntilAsyncQueriesDone() | Out-Null }} catch {{ }}
  $saved = $false
  for ($i=0; $i -lt 180; $i++) {{
    try {{
      $wb.Save() | Out-Null
      $saved = $true
      break
    }} catch {{
      Start-Sleep -Seconds 2
    }}
  }}
  if (-not $saved) {{ throw "Workbook save failed after refresh." }}
  $wb.Close($true) | Out-Null
}} finally {{
  try {{ $excel.Quit() | Out-Null }} catch {{ }}
  try {{ [System.Runtime.Interopservices.Marshal]::ReleaseComObject($excel) | Out-Null }} catch {{ }}
  [GC]::Collect()
  [GC]::WaitForPendingFinalizers()
}}
"""

    try:
        subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", ps],
            check=True,
            timeout=timeout_s,
        )
    except FileNotFoundError as exc:
        raise SystemExit("PowerShell not found; cannot auto-refresh workbook.") from exc
    except subprocess.TimeoutExpired as exc:
        raise SystemExit(f"Timed out refreshing workbook after {timeout_s}s.") from exc
    except subprocess.CalledProcessError as exc:
        raise SystemExit(f"Excel refresh failed (exit code {exc.returncode}).") from exc


def load_week(workbook: Path, sheet: str, header_row: int, week_posted: dt.date) -> WeekExtract:
    df = pd.read_excel(workbook, sheet_name=sheet, engine="openpyxl", header=header_row)
    required = {"Project", "Week Posted", "Expense Name", "Description", "Hours"}
    missing = sorted(required.difference(set(df.columns)))
    if missing:
        raise SystemExit(
            f"Missing expected columns on {sheet!r}: {', '.join(missing)}. "
            f"Found: {', '.join(map(str, df.columns))}"
        )

    df = df.copy()
    df["Week Posted"] = pd.to_datetime(df["Week Posted"], errors="coerce", format="mixed").dt.date

    df = df[df["Week Posted"] == week_posted]
    df = df[["Project", "Expense Name", "Description", "Hours"]]

    df["Project"] = df["Project"].map(_clean_text)
    df["Expense Name"] = df["Expense Name"].map(_clean_text)
    df["Description"] = df["Description"].map(_clean_text)
    df["Hours"] = pd.to_numeric(df["Hours"], errors="coerce").fillna(0.0)

    total_hours = float(df["Hours"].sum())
    blank_person_hours = float(df.loc[df["Expense Name"].isna(), "Hours"].sum())
    blank_project_hours = float(df.loc[df["Project"].isna(), "Hours"].sum())

    df = df[(df["Project"].notna()) & (df["Expense Name"].notna())]
    return WeekExtract(
        df=df,
        total_hours=total_hours,
        blank_person_hours=blank_person_hours,
        blank_project_hours=blank_project_hours,
    )


def load_hours_by_week(workbook: Path, sheet: str, header_row: int) -> pd.DataFrame:
    df = pd.read_excel(
        workbook,
        sheet_name=sheet,
        engine="openpyxl",
        header=header_row,
        usecols=["Week Posted", "Hours"],
    )
    df = df.copy()
    df["Week Posted"] = pd.to_datetime(df["Week Posted"], errors="coerce", format="mixed").dt.date
    df["Hours"] = pd.to_numeric(df["Hours"], errors="coerce").fillna(0.0)
    df = df.dropna(subset=["Week Posted"])

    df = df.groupby("Week Posted", as_index=False)["Hours"].sum().sort_values("Week Posted")
    df["Work Week Ending"] = df["Week Posted"].map(lambda d: d - dt.timedelta(days=7))
    return df[["Week Posted", "Work Week Ending", "Hours"]]


def trend_summary(hours_by_week: pd.DataFrame, selected_week_posted: dt.date) -> dict:
    if hours_by_week.empty:
        return {}

    series = hours_by_week.set_index("Week Posted")["Hours"]
    if selected_week_posted not in series.index:
        return {}

    this_hours = float(series.loc[selected_week_posted])
    prev_weeks = series.loc[series.index < selected_week_posted]
    prev_hours = float(prev_weeks.iloc[-1]) if len(prev_weeks) else None
    delta = (this_hours - prev_hours) if prev_hours is not None else None
    pct = (delta / prev_hours * 100.0) if prev_hours not in (None, 0.0) else None
    last4 = prev_weeks.tail(4)
    avg4 = float(last4.mean()) if len(last4) else None
    return {"this_hours": this_hours, "prev_hours": prev_hours, "delta": delta, "pct": pct, "avg4": avg4}


def export_trend_chart_png(hours_by_week: pd.DataFrame, out_png: Path, title: str, max_points: int) -> None:
    points = hours_by_week.tail(max_points).copy()
    if points.empty:
        raise SystemExit("No data available for chart.")

    labels = [d.strftime("%m/%d") for d in points["Work Week Ending"].tolist()]
    values = [float(h) for h in points["Hours"].tolist()]

    width, height = 980, 480
    pad_left, pad_right, pad_top, pad_bottom = 70, 30, 50, 70
    plot_w = width - pad_left - pad_right
    plot_h = height - pad_top - pad_bottom

    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default()

    # Title
    draw.text((pad_left, 15), title, fill="black", font=font)

    # Axes
    x0, y0 = pad_left, pad_top + plot_h
    x1, y1 = pad_left + plot_w, pad_top
    draw.line((x0, y0, x1, y0), fill="black", width=2)
    draw.line((x0, y0, x0, y1), fill="black", width=2)

    vmax = max(values) if values else 0.0
    vmax = vmax * 1.1 if vmax > 0 else 1.0
    vmin = 0.0

    # Y ticks
    ticks = 5
    for i in range(ticks + 1):
        t = i / ticks
        y = y0 - int(t * plot_h)
        val = vmin + t * (vmax - vmin)
        draw.line((x0 - 4, y, x0, y), fill="black", width=1)
        draw.text((10, y - 6), f"{val:.0f}", fill="black", font=font)

    if len(values) == 1:
        xs = [x0 + plot_w // 2]
    else:
        xs = [x0 + int(i * plot_w / (len(values) - 1)) for i in range(len(values))]

    def y_for(v: float) -> int:
        return y0 - int((v - vmin) / (vmax - vmin) * plot_h)

    pts = [(x, y_for(v)) for x, v in zip(xs, values)]

    # Line + markers
    for (ax, ay), (bx, by) in zip(pts, pts[1:]):
        draw.line((ax, ay, bx, by), fill="#1f77b4", width=3)
    for x, y in pts:
        r = 4
        draw.ellipse((x - r, y - r, x + r, y + r), fill="#1f77b4", outline="#1f77b4")

    # X labels (every N to avoid clutter)
    step = 1
    if len(labels) > 12:
        step = 2
    if len(labels) > 20:
        step = 3
    for i, (x, lbl) in enumerate(zip(xs, labels)):
        if i % step != 0 and i != len(labels) - 1:
            continue
        draw.line((x, y0, x, y0 + 4), fill="black", width=1)
        draw.text((x - 10, y0 + 10), lbl, fill="black", font=font)

    out_png.parent.mkdir(parents=True, exist_ok=True)
    img.save(out_png, format="PNG")


def _format_hours(value: float) -> str:
    return f"{value:.1f}".rstrip("0").rstrip(".")


def _top_description_bullets(descriptions: Iterable[str], max_items: int = 5) -> list[str]:
    counts = Counter([d for d in descriptions if d])
    if not counts:
        return []
    # Prefer most common; stable tie-breaker on text.
    items = sorted(counts.items(), key=lambda kv: (-kv[1], kv[0].lower()))
    bullets: list[str] = []
    for desc, n in items[:max_items]:
        suffix = f" (x{n})" if n > 1 else ""
        bullets.append(f"- {desc}{suffix}")
    return bullets


def render_markdown(
    extract: WeekExtract,
    window: WeekWindow,
    trend: dict | None = None,
    chart_path: Path | None = None,
) -> str:
    df = extract.df
    total_hours = extract.total_hours

    lines: list[str] = []
    lines.append(f"# Weekly Summary - Week Posted {window.week_posted.isoformat()}")
    lines.append("")
    lines.append(
        f"Worked period: {window.worked_start.isoformat()} (Sat) -> {window.worked_end.isoformat()} (Fri)"
    )
    lines.append("")
    lines.append(f"Total hours (all people): {_format_hours(total_hours)}")
    if extract.blank_person_hours or extract.blank_project_hours:
        lines.append(
            f"Note: {_format_hours(extract.blank_person_hours)}h has blank person and "
            f"{_format_hours(extract.blank_project_hours)}h has blank project in the source data."
        )
    lines.append("")

    if trend:
        lines.append("## Trend (Hours)")
        lines.append("")
        prev_hours = trend.get("prev_hours")
        delta = trend.get("delta")
        pct = trend.get("pct")
        avg4 = trend.get("avg4")
        if prev_hours is not None and delta is not None:
            sign = "+" if delta >= 0 else ""
            pct_text = f" ({sign}{pct:.1f}%)" if pct is not None else ""
            lines.append(
                f"- Vs prior week: {_format_hours(trend['this_hours'])} vs {_format_hours(prev_hours)} "
                f"({sign}{_format_hours(delta)}{pct_text})"
            )
        if avg4 is not None:
            sign = "+" if (trend['this_hours'] - avg4) >= 0 else ""
            lines.append(
                f"- Vs last 4-week avg: {_format_hours(trend['this_hours'])} vs {_format_hours(avg4)} "
                f"({sign}{_format_hours(trend['this_hours'] - avg4)})"
            )
        if chart_path is not None:
            lines.append(f"- Chart: {chart_path.as_posix()}")
        lines.append("")

    # Project/person summary.
    lines.append("## Project / Person Summary")
    lines.append("")

    grouped: dict[str, dict[str, list[str]]] = defaultdict(lambda: defaultdict(list))
    for project, person, desc in df[["Project", "Expense Name", "Description"]].itertuples(
        index=False, name=None
    ):
        if desc:
            grouped[project][person].append(desc)
        else:
            grouped[project][person]  # ensure key exists

    for project in sorted(grouped.keys(), key=str.lower):
        lines.append(f"- Project: {project}")
        people = grouped[project]
        for person in sorted(people.keys(), key=str.lower):
            bullets = _top_description_bullets(people[person])
            if bullets:
                lines.append(f"  - {person}:")
                lines.extend([f"    {b}" for b in bullets])
            else:
                lines.append(f"  - {person}: (no description provided)")
        lines.append("")

    # Hours table.
    lines.append("## Hours by Person")
    lines.append("")
    lines.append("| Person | Hours |")
    lines.append("|---|---:|")
    hours = (
        df.groupby("Expense Name", dropna=False)["Hours"]
        .sum()
        .sort_values(ascending=False)
        .reset_index()
    )
    for person, h in hours.itertuples(index=False, name=None):
        lines.append(f"| {person} | {_format_hours(float(h))} |")

    lines.append("")
    return "\n".join(lines)


def render_docx(
    extract: WeekExtract,
    window: WeekWindow,
    out_path: Path,
    trend: dict | None = None,
    chart_path: Path | None = None,
) -> None:
    df = extract.df
    doc = Document()
    doc.add_heading(f"Weekly Summary - Week Posted {window.week_posted.isoformat()}", level=1)
    doc.add_paragraph(
        f"Worked period: {window.worked_start.isoformat()} (Sat) -> {window.worked_end.isoformat()} (Fri)"
    )
    doc.add_paragraph(f"Total hours (all people): {_format_hours(extract.total_hours)}")
    if extract.blank_person_hours or extract.blank_project_hours:
        doc.add_paragraph(
            f"Note: {_format_hours(extract.blank_person_hours)}h has blank person and "
            f"{_format_hours(extract.blank_project_hours)}h has blank project in the source data."
        )

    if trend:
        doc.add_heading("Trend (Hours)", level=2)
        prev_hours = trend.get("prev_hours")
        delta = trend.get("delta")
        pct = trend.get("pct")
        avg4 = trend.get("avg4")
        if prev_hours is not None and delta is not None:
            sign = "+" if delta >= 0 else ""
            pct_text = f" ({sign}{pct:.1f}%)" if pct is not None else ""
            doc.add_paragraph(
                f"Vs prior week: {_format_hours(trend['this_hours'])} vs {_format_hours(prev_hours)} "
                f"({sign}{_format_hours(delta)}{pct_text})",
                style="List Bullet",
            )
        if avg4 is not None:
            sign = "+" if (trend["this_hours"] - avg4) >= 0 else ""
            doc.add_paragraph(
                f"Vs last 4-week avg: {_format_hours(trend['this_hours'])} vs {_format_hours(avg4)} "
                f"({sign}{_format_hours(trend['this_hours'] - avg4)})",
                style="List Bullet",
            )
        if chart_path is not None and chart_path.exists():
            doc.add_picture(str(chart_path))

    doc.add_heading("Hours by Person", level=2)
    hours = (
        df.groupby("Expense Name", dropna=False)["Hours"]
        .sum()
        .sort_values(ascending=False)
        .reset_index()
    )
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Person"
    table.rows[0].cells[1].text = "Hours"
    for person, h in hours.itertuples(index=False, name=None):
        row = table.add_row().cells
        row[0].text = str(person)
        row[1].text = _format_hours(float(h))

    doc.add_heading("Project / Person Summary", level=2)
    grouped: dict[str, dict[str, list[str]]] = defaultdict(lambda: defaultdict(list))
    for project, person, desc in df[["Project", "Expense Name", "Description"]].itertuples(
        index=False, name=None
    ):
        if desc:
            grouped[project][person].append(desc)
        else:
            grouped[project][person]

    for project in sorted(grouped.keys(), key=str.lower):
        doc.add_heading(f"Project: {project}", level=3)
        people = grouped[project]
        for person in sorted(people.keys(), key=str.lower):
            doc.add_paragraph(person, style="List Bullet")
            bullets = _top_description_bullets(people[person])
            if bullets:
                for b in bullets:
                    doc.add_paragraph(b[2:], style="List Bullet 2")
            else:
                doc.add_paragraph("(no description provided)", style="List Bullet 2")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Generate a weekly Markdown summary from the PICSummary workbook."
    )
    parser.add_argument(
        "--workbook",
        default="Vantagepoint In Excel Workbook v3.0.xlsm",
        help="Path to the Excel workbook (.xlsx/.xlsm).",
    )
    parser.add_argument(
        "--week-posted",
        default=None,
        help="Week Posted date to filter (YYYY-MM-DD). If omitted, uses the latest Week Posted in the sheet.",
    )
    parser.add_argument("--sheet", default=DEFAULT_SHEET, help="Sheet name to read.")
    parser.add_argument(
        "--header-row",
        type=int,
        default=None,
        help="0-based row index to use as the header row. If omitted, attempts to auto-detect.",
    )
    parser.add_argument(
        "--no-refresh",
        action="store_true",
        help="Skip workbook refresh (not recommended).",
    )
    parser.add_argument(
        "--format",
        choices=["md", "docx"],
        default="md",
        help="Output format.",
    )
    parser.add_argument(
        "--trend-weeks",
        type=int,
        default=12,
        help="Number of weeks to include in the trend chart.",
    )
    parser.add_argument(
        "--chart",
        action="store_true",
        help="Include a trend chart image.",
    )
    parser.add_argument(
        "--out",
        default=None,
        help="Write output to this path (defaults to stdout for md, or outputs/ for docx).",
    )
    args = parser.parse_args()

    workbook = Path(args.workbook)
    if not workbook.exists():
        raise SystemExit(f"Workbook not found: {str(workbook)!r}")

    header_row = args.header_row
    if header_row is None:
        header_row = detect_header_row(workbook, args.sheet)

    if not args.no_refresh:
        refresh_workbook_via_excel_com(workbook)

    week_posted = (
        parse_date(args.week_posted)
        if args.week_posted
        else detect_latest_week_posted(workbook, args.sheet, header_row)
    )
    window = week_window(week_posted)

    extract = load_week(workbook, args.sheet, header_row, week_posted)
    hours_by_week = load_hours_by_week(workbook, args.sheet, header_row)
    trend = trend_summary(hours_by_week, week_posted)

    chart_path: Path | None = None
    if args.chart:
        chart_path = Path("outputs") / f"trend-hours-{week_posted.isoformat()}.png"
        export_trend_chart_png(
            hours_by_week,
            chart_path,
            title="Billed Hours by Work Week",
            max_points=max(2, int(args.trend_weeks)),
        )

    if args.format == "docx":
        out_path = (
            Path(args.out)
            if args.out
            else Path("outputs") / f"weekly-summary-{week_posted.isoformat()}.docx"
        )
        render_docx(extract, window, out_path, trend=trend, chart_path=chart_path)
        return 0

    md = render_markdown(extract, window, trend=trend, chart_path=chart_path)

    if args.out:
        out_path = Path(args.out)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(md, encoding="utf-8")
    else:
        print(md)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
