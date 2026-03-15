#!/usr/bin/env python3
"""
Streamlit app to select equipment rows and generate a filtered Title V DOCX.
"""

from __future__ import annotations

import json
import os
import re
from glob import glob
from typing import Any, Dict, Iterable, List, Optional, Tuple

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DEFAULT_JSON_PATH = os.path.join("JSON", "title_v_permit_B2758_B2759.json")
TEMPLATE_PATH = os.path.join(
    "Title V permit", "B2758-59_Renewal_Rev6_Final_TitleV_09-29-2023.docx"
)
OUTPUT_PREFIX = "B2758_B2759_v"
OUTPUT_EXT = ".docx"


def load_json(path: str) -> Dict[str, Any]:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def find_section(sections: List[Dict[str, Any]], title: str) -> Optional[Dict[str, Any]]:
    title_up = title.strip().upper()
    for sec in sections:
        if (sec.get("title") or "").strip().upper() == title_up:
            return sec
        found = find_section(sec.get("sections", []), title)
        if found:
            return found
    return None


def get_equipment_tables(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    sec = find_section(data.get("sections", []), "EQUIPMENT")
    if not sec:
        return []
    return [item for item in sec.get("content", []) if item.get("type") == "table"]


def unit_match(units: Iterable[str], selected_units: set[str]) -> bool:
    return bool(set(units or []) & selected_units)


def unit_id_from_row(row: Dict[str, Any]) -> str:
    unit_id = row.get("unit_id", "") or ""
    if unit_id:
        return str(unit_id)
    units = row.get("applicable_units", []) or []
    return units[0] if units else ""


def dedupe_rows_by_unit(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    seen = set()
    out = []
    for row in rows:
        unit_id = unit_id_from_row(row)
        if not unit_id:
            out.append(row)
            continue
        if unit_id in seen:
            continue
        seen.add(unit_id)
        out.append(row)
    return out


def filter_table(table: Dict[str, Any], selected_units: set[str]) -> Optional[Dict[str, Any]]:
    rows = table.get("rows_as_objects", []) or []
    matched_rows = [
        row for row in rows if unit_match(row.get("applicable_units", []), selected_units)
    ]
    table_matches = unit_match(table.get("applicable_units", []), selected_units)
    if matched_rows:
        out = dict(table)
        out["rows_as_objects"] = matched_rows
        return out
    if table_matches:
        return dict(table)
    return None


def filter_section(
    section: Dict[str, Any], selected_units: set[str]
) -> Optional[Dict[str, Any]]:
    sec_units = section.get("applicable_units", [])
    sec_matches = unit_match(sec_units, selected_units)

    new_sec = dict(section)
    new_sec_sections = []
    if sec_matches:
        new_sec["content"] = list(section.get("content", []))
        for child in section.get("sections", []):
            filtered = filter_section(child, selected_units)
            if filtered:
                new_sec_sections.append(filtered)
        new_sec["sections"] = new_sec_sections
        return new_sec

    new_content = []
    for item in section.get("content", []):
        if item.get("type") == "table":
            filtered_table = filter_table(item, selected_units)
            if filtered_table:
                new_content.append(filtered_table)
            continue
        if unit_match(item.get("applicable_units", []), selected_units):
            new_content.append(item)

    for child in section.get("sections", []):
        filtered = filter_section(child, selected_units)
        if filtered:
            new_sec_sections.append(filtered)

    if not new_content and not new_sec_sections:
        return None

    new_sec["content"] = new_content
    new_sec["sections"] = new_sec_sections
    return new_sec


def clear_document(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def style_or_none(doc: Document, style_name: Optional[str]) -> Optional[str]:
    if not style_name:
        return None
    for style in doc.styles:
        if style.name == style_name:
            return style_name
    return None


def apply_paragraph_format(
    para,
    *,
    font_name: str = "Calibri",
    font_size: Optional[Pt] = None,
    bold: Optional[bool] = None,
    align: Optional[int] = None,
    spacing_before: Optional[Pt] = None,
    spacing_after: Optional[Pt] = None,
    line_spacing: Optional[float] = None,
    no_indent: bool = False,
) -> None:
    pf = para.paragraph_format
    if align is not None:
        para.alignment = align
    if spacing_before is not None:
        pf.space_before = spacing_before
    if spacing_after is not None:
        pf.space_after = spacing_after
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if no_indent:
        pf.left_indent = Pt(0)
        pf.first_line_indent = Pt(0)

    for run in para.runs:
        run.font.name = font_name
        if font_size is not None:
            run.font.size = font_size
        if bold is not None:
            run.font.bold = bold


def strip_paragraph_numbering(para) -> None:
    ppr = para._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is not None:
        ppr.remove(num_pr)


def add_paragraph(doc: Document, text: str, style_name: Optional[str]) -> Optional[Any]:
    if not text:
        return None
    para = doc.add_paragraph(text)
    if style_name:
        para.style = style_name
    apply_paragraph_format(para, font_name="Calibri")
    return para


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_set_text(
    cell,
    text: str,
    *,
    bold: Optional[bool] = None,
    align: Optional[int] = None,
    font_size: Optional[Pt] = None,
    spacing_before: Optional[Pt] = None,
    spacing_after: Optional[Pt] = None,
    line_spacing: Optional[float] = None,
) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.add_run(text)
    apply_paragraph_format(
        para,
        font_name="Calibri",
        font_size=font_size,
        bold=bold,
        align=align,
        spacing_before=spacing_before,
        spacing_after=spacing_after,
        line_spacing=line_spacing,
        no_indent=True,
    )


def split_table_title_lines(title: str) -> List[str]:
    return [line.strip() for line in title.splitlines() if line.strip()]


def ensure_table_description_style(doc: Document) -> Optional[str]:
    for style in doc.styles:
        if style.name == "Table Description":
            return style.name
    style = doc.styles.add_style("Table Description", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.bold = True
    pf = style.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return style.name


def add_table_title_block(doc: Document, title: str) -> None:
    lines = split_table_title_lines(title)
    if not lines:
        return

    table_desc_style = ensure_table_description_style(doc)

    for i, line in enumerate(lines):
        if i == 0 and line.upper().startswith("TABLE"):
            para = add_paragraph(doc, line, "Heading 4")
            if para:
                apply_paragraph_format(
                    para,
                    font_name="Calibri",
                    font_size=Pt(11),
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    spacing_after=Pt(6),
                    no_indent=True,
                )
                strip_paragraph_numbering(para)
            continue
        if "SOURCE-SPECIFIC APPLICABLE REQUIREMENTS" in line.upper():
            para = add_paragraph(doc, line, "Heading 4")
            if para:
                apply_paragraph_format(
                    para,
                    font_name="Calibri",
                    font_size=Pt(11),
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    spacing_after=Pt(6),
                    no_indent=True,
                )
                strip_paragraph_numbering(para)
            continue
        if line.upper().startswith("TANKS"):
            para = add_paragraph(doc, line, table_desc_style)
            if para:
                apply_paragraph_format(
                    para,
                    font_name="Calibri",
                    font_size=Pt(11),
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    spacing_before=Pt(2),
                    spacing_after=Pt(2),
                    line_spacing=1.0,
                    no_indent=True,
                )
            continue
        if line.upper().startswith("PLANT"):
            para = add_paragraph(doc, line, None)
            if para:
                apply_paragraph_format(
                    para,
                    font_name="Calibri",
                    font_size=Pt(11),
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    spacing_before=Pt(2),
                    spacing_after=Pt(2),
                    no_indent=True,
                )
            continue
        if line.upper().startswith("THE FOLLOWING") or line.upper().startswith("EACH OF THE FOLLOWING"):
            para = add_paragraph(doc, line, None)
            if para:
                apply_paragraph_format(
                    para,
                    font_name="Calibri",
                    font_size=Pt(11),
                    bold=None,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    spacing_before=Pt(2),
                    spacing_after=Pt(2),
                    no_indent=True,
                )
            continue

        para = add_paragraph(doc, line, None)
        if para:
            apply_paragraph_format(
                para,
                font_name="Calibri",
                font_size=Pt(11),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                spacing_before=Pt(2),
                spacing_after=Pt(2),
                no_indent=True,
            )


def write_table(doc: Document, table: Dict[str, Any]) -> None:
    headers = table.get("headers", []) or []
    rows = table.get("rows_as_objects", []) or []
    note_rows = table.get("note_rows", []) or []
    title = (table.get("title") or "").strip()
    if not headers:
        return

    if title:
        add_table_title_block(doc, title)

    tbl = doc.add_table(rows=0, cols=len(headers))
    tbl.style = "Table Grid"

    header_row = tbl.add_row()
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell_set_text(
            cell,
            h,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            font_size=Pt(11),
            spacing_before=Pt(2),
            spacing_after=Pt(2),
            line_spacing=1.0,
        )
        set_cell_shading(cell, "D9D9D9")

    for row in rows:
        data_row = tbl.add_row()
        for i, h in enumerate(headers):
            cell_set_text(
                data_row.cells[i],
                row.get(h, ""),
                bold=None,
                align=None,
                font_size=Pt(11),
                spacing_before=Pt(2),
                spacing_after=Pt(2),
                line_spacing=1.0,
            )

    for note in note_rows:
        texts = [c.strip() for c in note if c.strip()]
        if not texts:
            continue
        note_text = texts[0] if len(set(texts)) == 1 else " ".join(texts)
        row = tbl.add_row()
        cell_set_text(
            row.cells[0],
            note_text,
            bold=None,
            align=None,
            font_size=Pt(11),
            spacing_before=Pt(2),
            spacing_after=Pt(2),
            line_spacing=1.0,
        )
        for i in range(1, len(row.cells)):
            row.cells[0].merge(row.cells[i])
    doc.add_page_break()


def write_content_item(doc: Document, item: Dict[str, Any]) -> None:
    item_type = item.get("type")
    if item_type == "table":
        write_table(doc, item)
        return
    if item_type == "toc_entry":
        return

    text = item.get("text", "")
    style_name = style_or_none(doc, item.get("style"))
    if item_type == "list_item":
        num = item.get("num")
        if num is not None:
            text = f"{num}. {text}"
        para = add_paragraph(doc, text, style_name)
        if para:
            apply_paragraph_format(para, font_name="Calibri")
        return
    para = add_paragraph(doc, text, style_name)
    if para:
        apply_paragraph_format(para, font_name="Calibri")


def write_section(doc: Document, section: Dict[str, Any]) -> None:
    title = section.get("title", "")
    style_name = style_or_none(doc, section.get("style"))
    if not style_name and section.get("level"):
        style_name = style_or_none(doc, f"Heading {section.get('level')}")
    if title.strip().upper() == "EQUIPMENT" and (section.get("level") == 1):
        title = "II\tEQUIPMENT"
        style_name = "Heading 1"
    if title.strip().upper() == "SOURCE-SPECIFIC APPLICABLE REQUIREMENTS":
        title = "IV\tSOURCE-SPECIFIC APPLICABLE REQUIREMENTS"
        style_name = "Heading 1"
    para = add_paragraph(doc, title, style_name)
    if para:
        apply_paragraph_format(para, font_name="Calibri")
        if style_name == "Heading 1":
            strip_paragraph_numbering(para)

    for item in section.get("content", []):
        write_content_item(doc, item)
    for child in section.get("sections", []):
        write_section(doc, child)


def next_versioned_path(prefix: str, ext: str, base_dir: str = ".") -> str:
    pattern = re.compile(rf"^{re.escape(prefix)}(\d+){re.escape(ext)}$")
    versions = []
    for name in os.listdir(base_dir):
        m = pattern.match(name)
        if m:
            versions.append(int(m.group(1)))
    next_ver = max(versions) + 1 if versions else 1
    return os.path.join(base_dir, f"{prefix}{next_ver}{ext}")


def generate_docx(
    data: Dict[str, Any],
    selected_units: set[str],
    template_path: str,
    output_path: str,
) -> None:
    doc = Document(template_path)
    clear_document(doc)

    for sec in data.get("sections", []):
        filtered = filter_section(sec, selected_units)
        if filtered:
            write_section(doc, filtered)

    doc.save(output_path)


def main() -> None:
    st.set_page_config(page_title="Title V Equipment Selector", layout="wide")
    st.title("Title V Equipment Selector")

    json_files = sorted(glob(os.path.join("JSON", "*.json")))
    st.sidebar.header("Data Source")
    json_choice = st.sidebar.selectbox(
        "JSON file",
        options=json_files if json_files else [DEFAULT_JSON_PATH],
        index=0,
    )
    json_custom = st.sidebar.text_input("Or enter JSON path", value="")
    json_path = json_custom.strip() or json_choice

    if not os.path.exists(json_path):
        st.error(f"JSON not found: {json_path}")
        return
    if not os.path.exists(TEMPLATE_PATH):
        st.error(f"Template DOCX not found: {TEMPLATE_PATH}")
        return

    data = load_json(json_path)
    tables = get_equipment_tables(data)
    if not tables:
        st.error("No EQUIPMENT tables found in JSON.")
        return

    selected_units: set[str] = set()

    for idx, table in enumerate(tables, 1):
        title = (table.get("title") or "").strip() or f"EQUIPMENT Table {idx}"
        st.subheader(title)

        headers = table.get("headers", []) or []
        rows = table.get("rows_as_objects", []) or []
        rows = dedupe_rows_by_unit(rows)
        if not rows:
            st.caption("No rows in this table.")
            continue

        display_rows = [{h: r.get(h, "") for h in headers} for r in rows] if headers else rows
        df = pd.DataFrame(display_rows)
        df.insert(0, "Select", False)

        edited = st.data_editor(
            df,
            key=f"equipment_table_{idx}",
            hide_index=True,
            use_container_width=True,
            column_config={"Select": st.column_config.CheckboxColumn("Select")},
            disabled=[c for c in df.columns if c != "Select"],
        )

        chosen = edited[edited["Select"]].index.tolist()
        for row_idx in chosen:
            units = rows[row_idx].get("applicable_units", []) or []
            selected_units.update(units)

        st.caption(f"Selected units so far: {len(selected_units)}")

    if st.button("Generate DOCX", type="primary"):
        if not selected_units:
            st.warning("Select at least one row before generating.")
            return
        out_path = next_versioned_path(OUTPUT_PREFIX, OUTPUT_EXT)
        generate_docx(data, selected_units, TEMPLATE_PATH, out_path)
        st.success(f"Generated: {out_path}")


if __name__ == "__main__":
    main()
