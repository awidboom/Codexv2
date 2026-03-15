#!/usr/bin/env python3
"""
CLI tool to select equipment rows and generate a filtered Title V DOCX.

This is a Streamlit-free alternative to `titlev_streamlit_app.py`.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from typing import Any, Dict, Iterable, List, Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DEFAULT_JSON_PATH = os.path.join("JSON", "title_v_permit_B2758_B2759.json")
DEFAULT_TEMPLATE_PATH = os.path.join(
    "Title V permit", "B2758-59_Renewal_Rev6_Final_TitleV_09-29-2023.docx"
)
OUTPUT_PREFIX = "B2758_B2759_v"
OUTPUT_EXT = ".docx"


def load_json(path: str) -> Dict[str, Any]:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def find_section(sections: List[Dict[str, Any]], title: str) -> Optional[Dict[str, Any]]:
    title_up = title.strip().upper()
    for sec in sections:
        if (sec.get("title") or "").strip().upper() == title_up:
            return sec
        found = find_section(sec.get("sections", []), title)
        if found:
            return found
    return None


def get_equipment_tables(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    sec = find_section(data.get("sections", []), "EQUIPMENT")
    if not sec:
        return []
    return [item for item in sec.get("content", []) if item.get("type") == "table"]


def unit_match(units: Iterable[str], selected_units: set[str]) -> bool:
    return bool(set(units or []) & selected_units)


def unit_id_from_row(row: Dict[str, Any]) -> str:
    unit_id = row.get("unit_id", "") or ""
    if unit_id:
        return str(unit_id)
    units = row.get("applicable_units", []) or []
    return units[0] if units else ""


def dedupe_rows_by_unit(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    seen = set()
    out = []
    for row in rows:
        unit_id = unit_id_from_row(row)
        if not unit_id:
            out.append(row)
            continue
        if unit_id in seen:
            continue
        seen.add(unit_id)
        out.append(row)
    return out


def filter_table(table: Dict[str, Any], selected_units: set[str]) -> Optional[Dict[str, Any]]:
    rows = table.get("rows_as_objects", []) or []
    matched_rows = [
        row for row in rows if unit_match(row.get("applicable_units", []), selected_units)
    ]
    table_matches = unit_match(table.get("applicable_units", []), selected_units)
    if matched_rows:
        out = dict(table)
        out["rows_as_objects"] = matched_rows
        return out
    if table_matches:
        return dict(table)
    return None


def filter_section(section: Dict[str, Any], selected_units: set[str]) -> Optional[Dict[str, Any]]:
    sec_units = section.get("applicable_units", [])
    sec_matches = unit_match(sec_units, selected_units)

    new_sec = dict(section)
    new_sec_sections = []
    if sec_matches:
        new_sec["content"] = list(section.get("content", []))
        for child in section.get("sections", []):
            filtered = filter_section(child, selected_units)
            if filtered:
                new_sec_sections.append(filtered)
        new_sec["sections"] = new_sec_sections
        return new_sec

    new_content = []
    for item in section.get("content", []):
        if item.get("type") == "table":
            filtered_table = filter_table(item, selected_units)
            if filtered_table:
                new_content.append(filtered_table)
            continue
        if unit_match(item.get("applicable_units", []), selected_units):
            new_content.append(item)

    for child in section.get("sections", []):
        filtered = filter_section(child, selected_units)
        if filtered:
            new_sec_sections.append(filtered)

    if not new_content and not new_sec_sections:
        return None

    new_sec["content"] = new_content
    new_sec["sections"] = new_sec_sections
    return new_sec


def clear_document(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def style_or_none(doc: Document, style_name: Optional[str]) -> Optional[str]:
    if not style_name:
        return None
    for style in doc.styles:
        if style.name == style_name:
            return style_name
    return None


def apply_paragraph_format(
    para,
    *,
    font_name: str = "Calibri",
    font_size: Optional[Pt] = None,
    bold: Optional[bool] = None,
    align: Optional[int] = None,
    spacing_before: Optional[Pt] = None,
    spacing_after: Optional[Pt] = None,
    line_spacing: Optional[float] = None,
    no_indent: bool = False,
) -> None:
    pf = para.paragraph_format
    if align is not None:
        para.alignment = align
    if spacing_before is not None:
        pf.space_before = spacing_before
    if spacing_after is not None:
        pf.space_after = spacing_after
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if no_indent:
        pf.left_indent = Pt(0)
        pf.first_line_indent = Pt(0)

    for run in para.runs:
        run.font.name = font_name
        if font_size is not None:
            run.font.size = font_size
        if bold is not None:
            run.font.bold = bold


def strip_paragraph_numbering(para) -> None:
    ppr = para._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is not None:
        ppr.remove(num_pr)


def add_paragraph(doc: Document, text: str, style_name: Optional[str]) -> Optional[Any]:
    if not text:
        return None
    para = doc.add_paragraph(text)
    if style_name:
        para.style = style_name
    apply_paragraph_format(para, font_name="Calibri")
    return para


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_set_text(
    cell,
    text: str,
    *,
    bold: Optional[bool] = None,
    align: Optional[int] = None,
    font_size: Optional[Pt] = None,
    spacing_before: Optional[Pt] = None,
    spacing_after: Optional[Pt] = None,
    line_spacing: Optional[float] = None,
) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.add_run(text)
    apply_paragraph_format(
        para,
        font_name="Calibri",
        font_size=font_size,
        bold=bold,
        align=align,
        spacing_before=spacing_before,
        spacing_after=spacing_after,
        line_spacing=line_spacing,
        no_indent=True,
    )


def split_table_title_lines(title: str) -> List[str]:
    return [line.strip() for line in title.splitlines() if line.strip()]


def ensure_table_description_style(doc: Document) -> Optional[str]:
    for style in doc.styles:
        if style.name == "Table Description":
            return style.name
    style = doc.styles.add_style("Table Description", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.bold = True
    pf = style.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return style.name


def add_table_title_block(doc: Document, title: str) -> None:
    lines = split_table_title_lines(title)
    if not lines:
        return

    table_desc_style = ensure_table_description_style(doc)

    for i, line in enumerate(lines):
        if i == 0 and line.upper().startswith("TABLE"):
            para = add_paragraph(doc, line, "Heading 4")
            if para:
                apply_paragraph_format(
                    para,
                    font_name="Calibri",
                    font_size=Pt(11),
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    spacing_after=Pt(6),
                    no_indent=True,
                )
                strip_paragraph_numbering(para)
            continue
        if "SOURCE-SPECIFIC APPLICABLE REQUIREMENTS" in line.upper():
            para = add_paragraph(doc, line, "Heading 4")
            if para:
                apply_paragraph_format(
                    para,
                    font_name="Calibri",
                    font_size=Pt(11),
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    spacing_after=Pt(6),
                    no_indent=True,
                )
                strip_paragraph_numbering(para)
            continue
        if line.upper().startswith("TANKS"):
            para = add_paragraph(doc, line, table_desc_style)
            if para:
                apply_paragraph_format(
                    para,
                    font_name="Calibri",
                    font_size=Pt(11),
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    spacing_before=Pt(2),
                    spacing_after=Pt(2),
                    line_spacing=1.0,
                    no_indent=True,
                )
            continue
        if line.upper().startswith("PLANT"):
            para = add_paragraph(doc, line, None)
            if para:
                apply_paragraph_format(
                    para,
                    font_name="Calibri",
                    font_size=Pt(11),
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    spacing_before=Pt(2),
                    spacing_after=Pt(2),
                    no_indent=True,
                )
            continue
        if line.upper().startswith("THE FOLLOWING") or line.upper().startswith(
            "EACH OF THE FOLLOWING"
        ):
            para = add_paragraph(doc, line, None)
            if para:
                apply_paragraph_format(
                    para,
                    font_name="Calibri",
                    font_size=Pt(11),
                    bold=None,
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    spacing_before=Pt(2),
                    spacing_after=Pt(2),
                    no_indent=True,
                )
            continue

        para = add_paragraph(doc, line, None)
        if para:
            apply_paragraph_format(
                para,
                font_name="Calibri",
                font_size=Pt(11),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                spacing_before=Pt(2),
                spacing_after=Pt(2),
                no_indent=True,
            )


def write_table(doc: Document, table: Dict[str, Any]) -> None:
    headers = table.get("headers", []) or []
    rows = table.get("rows_as_objects", []) or []
    note_rows = table.get("note_rows", []) or []
    title = (table.get("title") or "").strip()
    if not headers:
        return

    if title:
        add_table_title_block(doc, title)

    tbl = doc.add_table(rows=0, cols=len(headers))
    tbl.style = "Table Grid"

    header_row = tbl.add_row()
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell_set_text(
            cell,
            h,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            font_size=Pt(11),
            spacing_before=Pt(2),
            spacing_after=Pt(2),
            line_spacing=1.0,
        )
        set_cell_shading(cell, "D9D9D9")

    for row in rows:
        data_row = tbl.add_row()
        for i, h in enumerate(headers):
            cell_set_text(
                data_row.cells[i],
                row.get(h, ""),
                bold=None,
                align=None,
                font_size=Pt(11),
                spacing_before=Pt(2),
                spacing_after=Pt(2),
                line_spacing=1.0,
            )

    for note in note_rows:
        texts = [c.strip() for c in note if c.strip()]
        if not texts:
            continue
        note_text = texts[0] if len(set(texts)) == 1 else " ".join(texts)
        row = tbl.add_row()
        cell_set_text(
            row.cells[0],
            note_text,
            bold=None,
            align=None,
            font_size=Pt(11),
            spacing_before=Pt(2),
            spacing_after=Pt(2),
            line_spacing=1.0,
        )
        for i in range(1, len(row.cells)):
            row.cells[0].merge(row.cells[i])
    doc.add_page_break()


def write_content_item(doc: Document, item: Dict[str, Any]) -> None:
    item_type = item.get("type")
    if item_type == "table":
        write_table(doc, item)
        return
    if item_type == "toc_entry":
        return

    text = item.get("text", "")
    style_name = style_or_none(doc, item.get("style"))
    if item_type == "list_item":
        num = item.get("num")
        if num is not None:
            text = f"{num}. {text}"
        para = add_paragraph(doc, text, style_name)
        if para:
            apply_paragraph_format(para, font_name="Calibri")
        return
    para = add_paragraph(doc, text, style_name)
    if para:
        apply_paragraph_format(para, font_name="Calibri")


def write_section(doc: Document, section: Dict[str, Any]) -> None:
    title = section.get("title", "")
    style_name = style_or_none(doc, section.get("style"))
    if not style_name and section.get("level"):
        style_name = style_or_none(doc, f"Heading {section.get('level')}")
    if title.strip().upper() == "EQUIPMENT" and (section.get("level") == 1):
        title = "II\tEQUIPMENT"
        style_name = "Heading 1"
    if title.strip().upper() == "SOURCE-SPECIFIC APPLICABLE REQUIREMENTS":
        title = "IV\tSOURCE-SPECIFIC APPLICABLE REQUIREMENTS"
        style_name = "Heading 1"
    para = add_paragraph(doc, title, style_name)
    if para:
        apply_paragraph_format(para, font_name="Calibri")
        if style_name == "Heading 1":
            strip_paragraph_numbering(para)

    for item in section.get("content", []):
        write_content_item(doc, item)
    for child in section.get("sections", []):
        write_section(doc, child)


def next_versioned_path(prefix: str, ext: str, base_dir: str = ".") -> str:
    pattern = re.compile(rf"^{re.escape(prefix)}(\d+){re.escape(ext)}$")
    versions = []
    for name in os.listdir(base_dir):
        m = pattern.match(name)
        if m:
            versions.append(int(m.group(1)))
    next_ver = max(versions) + 1 if versions else 1
    return os.path.join(base_dir, f"{prefix}{next_ver}{ext}")


def generate_docx(
    data: Dict[str, Any],
    selected_units: set[str],
    template_path: str,
    output_path: str,
) -> None:
    doc = Document(template_path)
    clear_document(doc)

    for sec in data.get("sections", []):
        filtered = filter_section(sec, selected_units)
        if filtered:
            write_section(doc, filtered)

    doc.save(output_path)


def normalize_unit_token(token: str) -> Optional[str]:
    t = (token or "").strip()
    if not t:
        return None
    t = t.replace("–", "-").replace("—", "-")
    t = re.sub(r"\s+", "", t)
    m = re.fullmatch(r"S-?\d+", t, flags=re.IGNORECASE)
    if m:
        digits = re.search(r"\d+", t)
        return f"S-{digits.group(0)}" if digits else None
    if t.isdigit():
        return f"S-{t}"
    m2 = re.search(r"\b(\d{2,5})\b", token)
    if m2 and re.search(r"tank", token, re.IGNORECASE):
        return f"S-{m2.group(1)}"
    return token.strip()


def parse_units(values: List[str]) -> List[str]:
    out: List[str] = []
    for v in values:
        parts = [p.strip() for p in (v or "").split(",") if p.strip()]
        for p in parts:
            n = normalize_unit_token(p)
            if n:
                out.append(n)
    return out


def list_equipment(data: Dict[str, Any]) -> int:
    tables = get_equipment_tables(data)
    if not tables:
        print("No EQUIPMENT tables found.", file=sys.stderr)
        return 2

    for idx, table in enumerate(tables, 1):
        title = (table.get("title") or "").strip() or f"EQUIPMENT Table {idx}"
        print(f"\n== {title} ==")
        rows = dedupe_rows_by_unit(table.get("rows_as_objects", []) or [])
        headers = table.get("headers", []) or []

        desc_key = "Description" if "Description" in headers else None
        for r in rows:
            unit_id = unit_id_from_row(r)
            desc = (r.get(desc_key) if desc_key else None) or ""
            print(f"{unit_id}\t{str(desc).strip()}")
    return 0


def resolve_units_from_equipment_tables(
    data: Dict[str, Any], requested: List[str], *, strict: bool
) -> set[str]:
    if not requested:
        raise SystemExit("No equipment provided. Use --equipment ... or --list-equipment.")

    requested_norm = [normalize_unit_token(x) for x in requested]
    requested_norm = [x for x in requested_norm if x]

    tables = get_equipment_tables(data)
    known_units: set[str] = set()
    for table in tables:
        for r in table.get("rows_as_objects", []) or []:
            unit_id = unit_id_from_row(r)
            if unit_id:
                known_units.add(unit_id)
            for u in (r.get("applicable_units", []) or []):
                known_units.add(str(u))

    selected = {u for u in requested_norm if u in known_units}
    missing = [u for u in requested_norm if u not in known_units]
    if missing and strict:
        raise SystemExit(
            "Unknown equipment id(s): "
            + ", ".join(missing)
            + ". Run with --list-equipment to see valid ids."
        )
    if missing:
        print(
            "Warning: unknown equipment id(s) ignored: " + ", ".join(missing),
            file=sys.stderr,
        )
    return selected


def build_arg_parser() -> argparse.ArgumentParser:
    ap = argparse.ArgumentParser(
        description="Generate a filtered Title V DOCX by selecting equipment ids (e.g., S-709)."
    )
    ap.add_argument("--json", dest="json_path", default=DEFAULT_JSON_PATH, help="Path to Title V JSON")
    ap.add_argument(
        "--template", dest="template_path", default=DEFAULT_TEMPLATE_PATH, help="Path to template DOCX"
    )
    ap.add_argument("--out", dest="out_path", default="", help="Output DOCX path")
    ap.add_argument(
        "--equipment",
        action="append",
        default=[],
        help="Equipment/unit id(s) to include (repeatable or comma-separated), e.g. S-709",
    )
    ap.add_argument(
        "--equipment-file",
        default="",
        help="Path to a text file containing one equipment id per line (or comma-separated).",
    )
    ap.add_argument("--list-equipment", action="store_true", help="List equipment ids from EQUIPMENT tables and exit")
    ap.add_argument(
        "--strict",
        action="store_true",
        help="Fail if any requested equipment id is not found in EQUIPMENT tables.",
    )
    return ap


def main(argv: Optional[List[str]] = None) -> int:
    args = build_arg_parser().parse_args(argv)

    if not os.path.exists(args.json_path):
        print(f"JSON not found: {args.json_path}", file=sys.stderr)
        return 2
    if not os.path.exists(args.template_path):
        print(f"Template DOCX not found: {args.template_path}", file=sys.stderr)
        return 2

    data = load_json(args.json_path)

    if args.list_equipment:
        return list_equipment(data)

    requested = parse_units(args.equipment or [])
    if args.equipment_file:
        with open(args.equipment_file, "r", encoding="utf-8") as f:
            for line in f:
                requested.extend(parse_units([line.strip()]))

    selected_units = resolve_units_from_equipment_tables(data, requested, strict=args.strict)
    if not selected_units:
        print("No matching equipment ids selected; nothing to generate.", file=sys.stderr)
        return 2

    out_path = args.out_path.strip() or next_versioned_path(OUTPUT_PREFIX, OUTPUT_EXT)
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    generate_docx(data, selected_units, args.template_path, out_path)
    print(out_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

