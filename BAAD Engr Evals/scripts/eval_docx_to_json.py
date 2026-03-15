import argparse
import json
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterator, List, Optional, Tuple


def _norm_space(s: str) -> str:
    return " ".join((s or "").replace("\xa0", " ").split()).strip()


def _unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    stem = path.stem
    suffix = path.suffix
    parent = path.parent
    i = 1
    while True:
        cand = parent / f"{stem}-{i}{suffix}"
        if not cand.exists():
            return cand
        i += 1


def _try_parse_date_str(text: str) -> Optional[str]:
    t = _norm_space(text)
    if not t:
        return None

    m = re.search(r"\b(\d{1,2})/(\d{1,2})/(\d{4})\b", t)
    if m:
        mm, dd, yyyy = m.group(1), m.group(2), m.group(3)
        return f"{int(mm)}/{int(dd)}/{yyyy}"

    months = (
        "January",
        "February",
        "March",
        "April",
        "May",
        "June",
        "July",
        "August",
        "September",
        "October",
        "November",
        "December",
    )
    m2 = re.search(
        rf"\b({'|'.join(months)})\s+(\d{{1,2}}),\s*(\d{{4}})\b",
        t,
        flags=re.IGNORECASE,
    )
    if m2:
        month, dd, yyyy = m2.group(1), int(m2.group(2)), m2.group(3)
        month = month[0].upper() + month[1:].lower()
        return f"{month} {dd}, {yyyy}"

    return None


def _table_to_markdown(table: Any) -> str:
    rows: List[List[str]] = []
    max_cols = 0
    for r in table.rows:
        cells = [_norm_space(c.text).replace("\n", "<br>") for c in r.cells]
        rows.append(cells)
        max_cols = max(max_cols, len(cells))

    if not rows or max_cols == 0:
        return ""

    # Normalize width
    for r in rows:
        if len(r) < max_cols:
            r.extend([""] * (max_cols - len(r)))

    header = rows[0]
    body = rows[1:] if len(rows) > 1 else []

    def esc(cell: str) -> str:
        return (cell or "").replace("|", "\\|")

    out: List[str] = []
    out.append("| " + " | ".join(esc(c) for c in header) + " |")
    out.append("|" + "|".join(["---"] * max_cols) + "|")
    for r in body:
        out.append("| " + " | ".join(esc(c) for c in r) + " |")
    return "\n".join(out)

def _iter_block_items(doc: Any) -> Iterator[Tuple[str, Any]]:
    # Yield ("paragraph", Paragraph) and ("table", Table) in document order, while
    # preserving the python-docx objects (so styles, runs, etc. resolve correctly).
    p_idx = 0
    t_idx = 0
    for child in doc.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            yield "paragraph", doc.paragraphs[p_idx]
            p_idx += 1
        elif tag == "tbl":
            yield "table", doc.tables[t_idx]
            t_idx += 1


SECTION_MAP = {
    "background": "background",
    "emissions calculation": "emission_calculations",
    "emission calculations": "emission_calculations",
    "emissions calculations": "emission_calculations",
    "cumulative increase": "cumulative_increase",
    "toxic screening analysis": "toxic_risk_screening_analysis",
    "best available control technology (bact)": "BACT",
    "best available control technology": "BACT",
    "bact": "BACT",
    "offsets": "offsets",
    "statement of compliance": "Statement_of_Compliance",
    "public notification": "public_notification",
    "public notice": "public_notification",
    "permit conditions": "conditions",
    "title v permit": "TitleV_permit",
    "recommendation": "recommendation",
}


def _normalize_heading_text(h: str) -> str:
    h = _norm_space(h)
    h = re.sub(r"^\s*\d+(\.\d+)?\s*[-–]?\s*", "", h)  # strip "1.0 -"
    return h.lower().strip()


def _heading_to_section_key(heading: str) -> Optional[str]:
    h = _normalize_heading_text(heading)
    for k, v in SECTION_MAP.items():
        if h.startswith(k):
            return v
    return None


@dataclass
class ExtractedMeta:
    doc_application_number: Optional[str]
    application_number: str
    plant_id: Optional[int]
    plant_name: str
    plant_address: str
    project_title: Optional[str]
    evaluation_date: Optional[str]


def _looks_like_heading(text: str) -> bool:
    t = _norm_space(text)
    if not t:
        return False
    # Common heading patterns in BAAD-style documents (sometimes styles are lost in conversions).
    if re.match(r"^\d+(\.\d+)?\s*[-–]\s*\S", t):
        return True
    if re.match(r"^[IVX]+\.\s*\d*\s*[-–]?\s*\S", t):
        return True
    # All-caps headings (avoid short acronyms)
    if len(t) >= 10 and t == t.upper() and re.search(r"[A-Z]{4,}", t):
        return True
    return False


def _extract_metadata(preamble_paras: List[str], all_paras: List[str], *, fallback_app: str) -> ExtractedMeta:
    doc_app: Optional[str] = None
    plant_id: Optional[int] = None
    plant_name = ""
    plant_address = ""
    project_title: Optional[str] = None
    evaluation_date: Optional[str] = None

    # plant name: first non-empty line after title, typically company line.
    for t in preamble_paras:
        if not t:
            continue
        if "engineering evaluation" in t.lower():
            continue
        if "air quality management district" in t.lower():
            continue
        if re.search(r"\bPlant Nos?\b", t, flags=re.IGNORECASE):
            continue
        if re.search(r"\bApplication No\.\b", t, flags=re.IGNORECASE):
            continue
        plant_name = t
        break

    # If we accidentally picked an address line, ignore it.
    if plant_name and (
        re.match(r"^\d{2,6}\s+\S+", plant_name)
        or re.search(r"\b(Street|St\.|Way|Avenue|Ave\.|Road|Rd\.|Boulevard|Blvd\.|San Francisco|Martinez)\b", plant_name, flags=re.IGNORECASE)
        or re.search(r"\bCA\s+\d{5}\b", plant_name)
    ):
        plant_name = ""

    # If the first line is a District header (e.g., Statement of Basis docs), try to find the permittee/facility name.
    if not plant_name or "air quality management district" in plant_name.lower():
        for t in preamble_paras:
            if re.search(r"\bTesoro\b", t, flags=re.IGNORECASE):
                plant_name = t
                break
        if not plant_name:
            for t in preamble_paras:
                if re.search(r"\bRefining\b", t, flags=re.IGNORECASE) and re.search(r"\bCompany\b", t, flags=re.IGNORECASE):
                    plant_name = t
                    break
    if not plant_name:
        for t in all_paras[:12]:
            if re.search(r"\b(refining|marketing|company|refinery)\b", t, flags=re.IGNORECASE) and not re.search(r"\bApplication No\.\b", t, flags=re.IGNORECASE):
                plant_name = t
                break

    # Plant Nos line -> first plant id
    for t in preamble_paras:
        m = re.search(r"\bPlant Nos?\.\s*([0-9,\s]+)", t, flags=re.IGNORECASE)
        if m:
            nums = re.findall(r"\d{3,6}", m.group(1))
            if nums:
                plant_id = int(nums[0])
            break

    for idx, t in enumerate(preamble_paras):
        if re.match(r"^\d{1,6}\s+\S+", t) and re.search(r"\b(Street|St\.|Way|Avenue|Ave\.|Road|Rd\.|Boulevard|Blvd\.|Drive|Dr\.|Lane|Ln\.|Court|Ct\.)\b", t, flags=re.IGNORECASE):
            plant_address = t
            if idx + 1 < len(preamble_paras) and re.search(r"\b[A-Za-z .'-]+,\s*[A-Z]{2}\s+\d{5}\b", preamble_paras[idx + 1]):
                plant_address = f"{plant_address}, {preamble_paras[idx + 1]}"
            break
    if not plant_address:
        for t in all_paras[:12]:
            if re.match(r"^\d{1,6}\s+\S+", t) and re.search(r"\b(Street|St\.|Way|Avenue|Ave\.|Road|Rd\.|Boulevard|Blvd\.|Drive|Dr\.|Lane|Ln\.|Court|Ct\.)\b", t, flags=re.IGNORECASE):
                plant_address = t
                break

    # Application No line
    for t in preamble_paras:
        m = re.search(r"\bApplication No\.\s*([0-9]{1,6})\b", t, flags=re.IGNORECASE)
        if m:
            doc_app = m.group(1)
            break

    # Project title from intro sentence if present
    for t in preamble_paras:
        m = re.search(r"for\s+(?:the\s+)?[^.]*?\bapplication\b.*?\bfor\b.*?\bfor\s+its\s+(.+?),\s*Application No\.", t, flags=re.IGNORECASE)
        if m:
            project_title = _norm_space(m.group(1))
            break
        m2 = re.search(r"for\s+its\s+(.+?),\s*Application No\.", t, flags=re.IGNORECASE)
        if m2:
            project_title = _norm_space(m2.group(1))
            break

    # Evaluation date: prefer explicit "Date:" near end
    for t in reversed(all_paras[-300:]):
        if not t:
            continue
        if re.search(r"\bDate\s*:\s*", t, flags=re.IGNORECASE):
            ds = _try_parse_date_str(t)
            if ds:
                evaluation_date = ds
                break

    if not evaluation_date:
        # last date-like occurrence in the whole document
        for t in reversed(all_paras):
            ds = _try_parse_date_str(t)
            if ds:
                evaluation_date = ds
                break

    if not project_title:
        # Common for SOB-style docs: project title appears as centered multi-line preamble.
        for t in preamble_paras:
            if re.search(r"\bMajor Facility Review Permit\b", t, flags=re.IGNORECASE):
                project_title = t
                break
        if not project_title:
            for t in preamble_paras:
                if re.search(r"\bStatement of Basis\b", t, flags=re.IGNORECASE):
                    project_title = t
                    break

    app_number = fallback_app
    if doc_app:
        app_number = fallback_app

    return ExtractedMeta(
        doc_application_number=doc_app,
        application_number=app_number,
        plant_id=plant_id,
        plant_name=plant_name,
        plant_address=plant_address,
        project_title=project_title,
        evaluation_date=evaluation_date,
    )


def docx_eval_to_json(
    docx_path: Path,
    *,
    application_number: str,
) -> Dict[str, Any]:
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        raise RuntimeError("python-docx is required (pip install python-docx)") from e

    doc = Document(str(docx_path))

    preamble: List[str] = []
    section_buffers: Dict[str, List[str]] = {}
    raw_blocks: List[Dict[str, Any]] = []

    current_heading: Optional[str] = None
    current_key: Optional[str] = None
    mapped_text_chars = 0

    # For equipment extraction: detect first "Sources Covered By Application" table.
    equipment: Dict[str, str] = {}
    last_para_text = ""
    saw_sources_table = False

    all_para_texts: List[str] = []

    for kind, item in _iter_block_items(doc):
        if kind == "paragraph":
            txt = _norm_space(item.text)
            if txt:
                all_para_texts.append(txt)
            style = item.style.name if getattr(item, "style", None) else ""

            is_h1 = style == "Heading 1"
            is_heading = (is_h1 or _looks_like_heading(txt)) and bool(txt)
            if is_heading and txt:
                current_heading = txt
                current_key = _heading_to_section_key(txt)
                if current_key:
                    section_buffers.setdefault(current_key, [])
                else:
                    raw_blocks.append({"heading": txt, "text": ""})
                last_para_text = txt
                continue

            # Preamble: before the first Heading 1 section
            if current_heading is None:
                if txt:
                    preamble.append(txt)
                last_para_text = txt
                continue

            if not txt:
                last_para_text = ""
                continue

            if current_key:
                section_buffers[current_key].append(txt)
                mapped_text_chars += len(txt)
            else:
                # Append into latest raw block if it matches current heading
                if raw_blocks and raw_blocks[-1].get("heading") == current_heading:
                    raw_blocks[-1]["text"] = (raw_blocks[-1].get("text") or "").rstrip() + "\n\n" + txt
            last_para_text = txt

        elif kind == "table":
            md = _table_to_markdown(item)
            if not md.strip():
                continue

            # Equipment table heuristic: caption paragraph immediately preceding.
            if (
                not saw_sources_table
                and last_para_text
                and re.search(r"Table\s*1\b", last_para_text, flags=re.IGNORECASE)
                and re.search(r"Sources\s+Covered\s+By\s+Application", last_para_text, flags=re.IGNORECASE)
            ):
                # Expect first column contains source ids like S-115, S-126, etc.
                for r in item.rows[1:]:
                    cells = [_norm_space(c.text) for c in r.cells]
                    if not cells:
                        continue
                    sid = cells[0].splitlines()[0].strip()
                    if re.match(r"^[SAF]-\d{2,6}\b", sid):
                        desc = cells[1] if len(cells) > 1 else ""
                        equipment[sid] = _norm_space(desc)
                saw_sources_table = True

            if current_heading is None:
                preamble.append(md)
            elif current_key:
                section_buffers[current_key].append(md)
                mapped_text_chars += len(md)
            else:
                if raw_blocks and raw_blocks[-1].get("heading") == current_heading:
                    raw_blocks[-1]["text"] = (raw_blocks[-1].get("text") or "").rstrip() + "\n\n" + md

    meta = _extract_metadata(preamble, all_para_texts, fallback_app=application_number)

    now = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")
    out: Dict[str, Any] = {
        "schema_version": "v1.2-barr-custom",
        "source": {
            "type": "docx",
            "created_local": now,
            "docx": str(docx_path.as_posix()),
            "doc_application_number": meta.doc_application_number,
        },
        "application_number": meta.application_number,
        "plant": {"name": meta.plant_name or "", "plant_id": meta.plant_id, "address": meta.plant_address or ""},
        "evaluation_date": meta.evaluation_date,
        "permit_date": None,
        "project_title": meta.project_title,
        "equipment": [{"id": k, "description": v} for k, v in sorted(equipment.items())],
        "background": {"text": "\n\n".join(section_buffers.get("background", [])).strip()},
        "emission_calculations": {"text": "\n\n".join(section_buffers.get("emission_calculations", [])).strip()},
        "cumulative_increase": {"text": "\n\n".join(section_buffers.get("cumulative_increase", [])).strip()},
        "emissions": {
            "common_fugitives": {
                "narrative": "",
                "predicted_component_counts_table": {
                    "columns": ["Component Type", "Service", "Predicted Count"],
                    "rows": [],
                },
                "permitted_fugitive_totals": {
                    "units": {"lb_per_day": "lb/day", "tpy": "tons/year"},
                    "POC": {"lb_per_day": None, "tpy": None},
                },
            },
            "by_equipment": {},
        },
        "toxic_risk_screening_analysis": {
            "status": "unknown",
            "narrative": "\n\n".join(section_buffers.get("toxic_risk_screening_analysis", [])).strip(),
        },
        "BACT": {"text": "\n\n".join(section_buffers.get("BACT", [])).strip()},
        "offsets": {
            "narrative": "\n\n".join(section_buffers.get("offsets", [])).strip(),
            "table": {
                "columns": ["Pollutant", "Project Increase (tpy)", "Offset Ratio", "Required Offsets (tpy)"],
                "rows": [],
            },
        },
        "PSD_applicability": {"narrative": ""},
        "CEQA": {"narrative": ""},
        "Statement_of_Compliance": "\n\n".join(section_buffers.get("Statement_of_Compliance", [])).strip(),
        "public_notification": {"text": "\n\n".join(section_buffers.get("public_notification", [])).strip()},
        "conditions": {"text": "\n\n".join(section_buffers.get("conditions", [])).strip()},
        "permit_conditions": {"condition_number": None, "items": []},
        "TitleV_permit": {
            "narrative": "\n\n".join(section_buffers.get("TitleV_permit", [])).strip(),
            "revisions": [],
            "affected_sources": [],
        },
        "recommendation": {"text": "\n\n".join(section_buffers.get("recommendation", [])).strip()},
        "raw_sections": {
            "preamble": "\n\n".join(preamble).strip(),
            "blocks": raw_blocks,
            "unmapped_blocks": [],
        },
    }

    # Fallback: if we didn't confidently segment any sections (common in .doc->.docx conversions),
    # ensure the document's text is still available for RAG by stuffing it into Background.
    if mapped_text_chars < 2000 and not out["background"]["text"]:
        pre = out.get("raw_sections", {}).get("preamble") or ""
        if isinstance(pre, str) and len(pre) > 2000:
            out["background"]["text"] = pre

    return out


def main(argv: Optional[List[str]] = None) -> int:
    ap = argparse.ArgumentParser(description="Convert an Engineering Evaluation DOCX into v1.2-barr-custom JSON.")
    ap.add_argument("--docx", required=True, help="Input .docx path")
    ap.add_argument(
        "--application-number",
        required=True,
        help="Application number to use in output JSON (use a unique id for templates).",
    )
    ap.add_argument("--out", required=True, help="Output .json path")
    args = ap.parse_args(argv)

    docx_path = Path(args.docx)
    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    obj = docx_eval_to_json(docx_path, application_number=str(args.application_number))
    out_path = _unique_path(out_path)
    out_path.write_text(json.dumps(obj, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(str(out_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
