import argparse
import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.shared import Pt  # type: ignore


def _norm_space(s: str) -> str:
    return " ".join((s or "").replace("\xa0", " ").split()).strip()


def _read_json(path: Path) -> Dict[str, Any]:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


SECTION_SPECS: List[Tuple[str, str]] = [
    ("background", "1.0 - Background"),
    ("emission_calculations", "2.0 Emissions Calculation"),
    ("cumulative_increase", "3.0 Cumulative Increase"),
    ("toxic_risk_screening_analysis", "4.0 Toxic Screening Analysis"),
    ("BACT", "5.0 Best Available Control Technology (BACT)"),
    ("offsets", "6.0 Offsets"),
    ("Statement_of_Compliance", "7.0 Statement of Compliance"),
    ("public_notification", "8.0 Public Notification"),
    ("conditions", "9.0 Permit Conditions"),
    ("TitleV_permit", "10.0 Title V Permit"),
    ("recommendation", "11.0 Recommendation"),
]


def _apply_run_format(
    run: Any,
    *,
    font_name: str = "Calibri",
    font_size: int = 10,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    run.underline = underline


def _format_paragraph(
    paragraph: Any,
    *,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
    font_name: str = "Calibri",
    font_size: int = 10,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    for run in paragraph.runs:
        _apply_run_format(
            run,
            font_name=font_name,
            font_size=font_size,
            bold=bold,
            italic=italic,
            underline=underline,
        )


def _add_formatted_paragraph(
    doc: Any,
    text: str,
    *,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
    font_size: int = 10,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
) -> Any:
    paragraph = doc.add_paragraph()
    paragraph.add_run(text)
    _format_paragraph(
        paragraph,
        align=align,
        font_size=font_size,
        bold=bold,
        italic=italic,
        underline=underline,
    )
    return paragraph


def _add_blank_paragraph(doc: Any) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)


def _format_table_cell_paragraph(paragraph: Any, *, bold: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    for run in paragraph.runs:
        _apply_run_format(run, font_size=10, bold=bold)


def _clear_doc_body(doc: Any) -> None:
    body = doc._element.body  # type: ignore[attr-defined]
    for child in list(body):
        # Keep section properties
        if child.tag.rsplit("}", 1)[-1] == "sectPr":
            continue
        body.remove(child)


def _split_markdown_row(line: str) -> List[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _is_md_table_block(lines: List[str]) -> bool:
    if len(lines) < 2:
        return False
    if not all(l.strip().startswith("|") for l in lines[:2]):
        return False
    sep = lines[1].strip()
    return bool(re.match(r"^\|?[\s:\-]+\|[\s:\-\|]+\|?$", sep))


def _extract_md_table_blocks(text: str) -> List[Tuple[str, Optional[List[str]]]]:
    # Returns ordered blocks: ("table", [lines]) or ("text", [lines]) as a single joined string.
    raw_lines = text.splitlines()
    blocks: List[Tuple[str, List[str]]] = []
    buf: List[str] = []
    in_table = False
    table_buf: List[str] = []

    def flush_text():
        nonlocal buf
        if buf:
            blocks.append(("text", buf))
            buf = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            blocks.append(("table", table_buf))
            table_buf = []

    for line in raw_lines:
        if line.strip().startswith("|"):
            if not in_table:
                # Starting a table: flush any text collected so far.
                flush_text()
                in_table = True
            table_buf.append(line.rstrip())
        else:
            if in_table:
                # End table on first non-table line (even blank)
                flush_table()
                in_table = False
            buf.append(line.rstrip())

    if in_table:
        flush_table()
    flush_text()

    out: List[Tuple[str, Optional[List[str]]]] = []
    for kind, lines in blocks:
        if kind == "table" and _is_md_table_block(lines):
            out.append(("table", lines))
        else:
            out.append(("text", ["\n".join(lines).strip()]))
    return out


def _paragraph_kind(text: str) -> str:
    t = text.strip()
    if not t:
        return "blank"
    if re.match(r"^Table\s+\d+\s*[-–]", t, flags=re.IGNORECASE):
        return "table_caption"
    if re.match(r"^\d+\.\d+\.\d+\s*[-–]", t):
        return "level3"
    if re.match(r"^\d+\.\d+\s*[-–]", t):
        return "level2"
    if re.match(r"^Section\s+\d", t):
        return "level3"
    if t.endswith(":") and len(t) <= 100:
        return "level2"
    return "normal"


def _add_paragraphs(doc: Any, text: str) -> None:
    for para_text in [p for p in (text or "").split("\n\n") if p.strip()]:
        kind = _paragraph_kind(para_text)
        if kind == "table_caption":
            _add_formatted_paragraph(
                doc,
                para_text.strip(),
                align=WD_ALIGN_PARAGRAPH.CENTER,
                font_size=10,
                bold=True,
            )
        elif kind == "level2":
            _add_formatted_paragraph(
                doc,
                para_text.strip(),
                align=WD_ALIGN_PARAGRAPH.LEFT,
                font_size=10,
                bold=True,
                italic=True,
            )
        elif kind == "level3":
            _add_formatted_paragraph(
                doc,
                para_text.strip(),
                align=WD_ALIGN_PARAGRAPH.LEFT,
                font_size=10,
                italic=True,
                underline=True,
            )
        else:
            _add_formatted_paragraph(
                doc,
                para_text.strip(),
                align=WD_ALIGN_PARAGRAPH.LEFT,
                font_size=10,
            )
        _add_blank_paragraph(doc)


def _add_markdown_table(doc: Any, lines: List[str]) -> None:
    rows = [_split_markdown_row(l) for l in lines if l.strip()]
    if len(rows) < 2:
        return
    # Drop the separator row (row 1)
    header = rows[0]
    body = [r for i, r in enumerate(rows[2:], start=2)]
    cols = max(len(header), *(len(r) for r in body)) if body else len(header)

    def pad(r: List[str]) -> List[str]:
        if len(r) < cols:
            r = r + [""] * (cols - len(r))
        return r[:cols]

    header = pad(header)
    body = [pad(r) for r in body]

    tbl = doc.add_table(rows=1 + len(body), cols=cols)
    try:
        tbl.style = "Table Grid"
    except Exception:
        pass

    for j, cell_text in enumerate(header):
        tbl.cell(0, j).text = cell_text.replace("<br>", "\n")
        for paragraph in tbl.cell(0, j).paragraphs:
            _format_table_cell_paragraph(paragraph, bold=True)
    for i, r in enumerate(body, start=1):
        for j, cell_text in enumerate(r):
            tbl.cell(i, j).text = cell_text.replace("<br>", "\n")
            for paragraph in tbl.cell(i, j).paragraphs:
                _format_table_cell_paragraph(paragraph, bold=False)


def _section_text(eval_obj: Dict[str, Any], section_key: str) -> str:
    if section_key == "background":
        return (eval_obj.get("background") or {}).get("text") or ""
    if section_key == "emission_calculations":
        return (eval_obj.get("emission_calculations") or {}).get("text") or ""
    if section_key == "cumulative_increase":
        return (eval_obj.get("cumulative_increase") or {}).get("text") or ""
    if section_key == "toxic_risk_screening_analysis":
        return (eval_obj.get("toxic_risk_screening_analysis") or {}).get("narrative") or ""
    if section_key == "BACT":
        return (eval_obj.get("BACT") or {}).get("text") or ""
    if section_key == "offsets":
        return (eval_obj.get("offsets") or {}).get("narrative") or ""
    if section_key == "Statement_of_Compliance":
        return str(eval_obj.get("Statement_of_Compliance") or "")
    if section_key == "public_notification":
        return (eval_obj.get("public_notification") or {}).get("text") or ""
    if section_key == "conditions":
        return (eval_obj.get("conditions") or {}).get("text") or ""
    if section_key == "TitleV_permit":
        return (eval_obj.get("TitleV_permit") or {}).get("narrative") or ""
    if section_key == "recommendation":
        return (eval_obj.get("recommendation") or {}).get("text") or ""
    return ""


def eval_json_to_docx(
    eval_obj: Dict[str, Any],
    *,
    template_docx: Path,
    out_docx: Path,
    include_sections: Optional[List[str]] = None,
) -> None:
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        raise RuntimeError("python-docx is required (pip install python-docx)") from e

    doc = Document(str(template_docx))
    _clear_doc_body(doc)

    app = _norm_space(str(eval_obj.get("application_number") or ""))
    plant = eval_obj.get("plant") or {}
    plant_name = _norm_space(str(plant.get("name") or ""))
    plant_id = plant.get("plant_id")
    plant_address = _norm_space(str(plant.get("address") or eval_obj.get("address") or ""))
    project_title = _norm_space(str(eval_obj.get("project_title") or ""))
    eval_date = eval_obj.get("evaluation_date")

    _add_formatted_paragraph(
        doc,
        "DRAFT ENGINEERING EVALUATION",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        font_size=12,
        bold=True,
    )
    if plant_name:
        _add_formatted_paragraph(
            doc,
            plant_name,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            font_size=10,
            bold=True,
        )
    if plant_address:
        _add_formatted_paragraph(
            doc,
            plant_address,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            font_size=10,
            bold=True,
        )
    if plant_id is not None:
        _add_formatted_paragraph(
            doc,
            f"Plant No.: {plant_id}",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            font_size=10,
            bold=True,
        )
    if app:
        _add_formatted_paragraph(
            doc,
            f"Application No.: {app}",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            font_size=10,
            bold=True,
        )
    _add_blank_paragraph(doc)

    if plant_name and project_title and app:
        intro = (
            f"This document is the Air District’s Engineering Evaluation Report for the {plant_name} "
            f"application for an Authority to Construct for its {project_title}, Application No. {app}."
        )
        _add_formatted_paragraph(
            doc,
            intro,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
        )
        _add_blank_paragraph(doc)
    if eval_date:
        _add_formatted_paragraph(
            doc,
            f"Evaluation Date: {eval_date}",
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
        )
        _add_blank_paragraph(doc)

    selected = {s.strip() for s in (include_sections or []) if s.strip()}
    section_specs: List[Tuple[str, str]] = []
    for section_key, heading in SECTION_SPECS:
        if selected and section_key not in selected:
            continue
        section_specs.append((heading, _section_text(eval_obj, section_key)))

    for heading, content in section_specs:
        _add_formatted_paragraph(
            doc,
            heading,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font_size=10,
            bold=True,
            underline=True,
        )
        content = str(content or "").strip()
        if not content:
            _add_formatted_paragraph(
                doc,
                "(empty)",
                align=WD_ALIGN_PARAGRAPH.LEFT,
                font_size=10,
            )
            _add_blank_paragraph(doc)
            continue

        blocks = _extract_md_table_blocks(content)
        for kind, payload in blocks:
            if kind == "table" and payload:
                _add_markdown_table(doc, payload)
                _add_blank_paragraph(doc)
            else:
                txt = (payload[0] if payload else "").strip()
                if txt:
                    _add_paragraphs(doc, txt)

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


def main(argv: Optional[List[str]] = None) -> int:
    ap = argparse.ArgumentParser(description="Render a v1.2-barr-custom evaluation JSON to a styled DOCX using a template.")
    ap.add_argument("--eval-json", required=True, help="Input evaluation JSON path")
    ap.add_argument(
        "--template-docx",
        default="Engineering Evaluation v5 7-18-2022.docx",
        help="Template DOCX (default: Engineering Evaluation v5 7-18-2022.docx)",
    )
    ap.add_argument(
        "--sections",
        default="",
        help="Comma-separated section keys to include (e.g. background,recommendation). Default renders all sections.",
    )
    ap.add_argument("--out-docx", required=True, help="Output DOCX path")
    args = ap.parse_args(argv)

    eval_path = Path(args.eval_json)
    template_path = Path(args.template_docx)
    out_path = Path(args.out_docx)
    include_sections = [s.strip() for s in str(args.sections or "").split(",") if s.strip()]

    obj = _read_json(eval_path)
    eval_json_to_docx(
        obj,
        template_docx=template_path,
        out_docx=out_path,
        include_sections=include_sections,
    )
    print(str(out_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
